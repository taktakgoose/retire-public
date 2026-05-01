#!/usr/bin/env python3
"""
retirement_to_docx.py — Year-by-year retirement action plan.

Reads retirement.csv (output of the C simulation) and produces
retirement_plan.docx with one section per simulated year showing:
  • Actions Required       — which accounts to draw from and how much
  • Expected Income        — CPP, OAS, rent, salary, and other passive sources
  • Portfolio Management   — TFSA room, surplus investing, and non-reg tax harvesting
  • Tax Summary            — gross income, pension split, capital gains, and taxes
  • Tax Filing Notes       — checklist of forms, elections, and credits to claim
  • Year-End Position      — cash reserve and total net worth

Usage:
    python3 retirement_to_docx.py [input.csv] [output.docx]

Defaults to retirement.csv -> retirement_plan.docx in the current directory.
"""

import sys
import csv
import os
import re
import math
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Column indices (0-based) — fixed summary block (cols 0–43)
# ---------------------------------------------------------------------------
C_YEAR     = 0;  C_CASH     = 1
C_EXPENSES = 4

# Ralph per-account summary (5-14)
C_CH_RRSP    = 5;  C_CH_DCPP    = 6;  C_CH_NONREG  = 7;  C_CH_TFSA    = 8
C_CH_RENT    = 11; C_CH_OAS     = 12; C_CH_CPP     = 13; C_CH_POLARON = 14

# Sarah per-account summary (15-24)
C_LI_RRSP    = 15; C_LI_DCPP    = 16; C_LI_NONREG  = 17; C_LI_TFSA    = 18
C_LI_RENT    = 21; C_LI_OAS     = 22; C_LI_CPP     = 23; C_LI_POLARON = 24

C_GROSS_NW   = 25   # Gross net worth (no tax deductions)
C_NETWORTH   = 26   # After-tax net worth (deemed disposition taxes applied)

# Ralph tax detail (27-36)
C_CH_AGE     = 27; C_CH_SALARY  = 28
C_CH_INCOME  = 30; C_CH_REG     = 31; C_CH_CAPITAL = 32
C_CH_TAKE    = 33; C_CH_TAX     = 34; C_CH_RATE    = 35
C_CH_HARVEST = 36  # Notional sell value for ACB-reset harvest

# Sarah tax detail (37-47)
C_LI_AGE     = 37; C_LI_SALARY  = 38
C_LI_INCOME  = 40; C_LI_REG     = 41; C_LI_CAPITAL = 42
C_LI_TAKE    = 43; C_LI_TAX     = 44; C_LI_RATE    = 45
C_LI_HARVEST = 46  # Notional sell value for ACB-reset harvest

# ---------------------------------------------------------------------------
# Column indices — asset detail block (starts at col 47)
# Layout (fixed order from init_years):
#   8 regular assets × 5 cols each  (RRSP×2, DCPP×2, NonReg×2, TFSA×2)
#   4 property assets × 8 cols each (Robertson×2, Moonbrook×2)
#   8 regular assets × 5 cols each  (Rent×2, OAS×2, CPP×2, Polaron×2)
# Each regular:  Name, Value, Debit, Room, Growth
# Each property: Name, Value, Debit, Room, Growth, MtgBal, MtgInt, MtgPrinc
# ---------------------------------------------------------------------------
_AB = 47                               # asset-block base column (0-based)

# Regular group 1 (5 cols each)
A_CH_RRSP_VAL  = _AB +  1;  A_CH_RRSP_DEB  = _AB +  2
A_LI_RRSP_VAL  = _AB +  6;  A_LI_RRSP_DEB  = _AB +  7
A_CH_DCPP_VAL  = _AB + 11;  A_CH_DCPP_DEB  = _AB + 12
A_LI_DCPP_VAL  = _AB + 16;  A_LI_DCPP_DEB  = _AB + 17
A_CH_NREG_VAL  = _AB + 21;  A_CH_NREG_DEB  = _AB + 22
A_LI_NREG_VAL  = _AB + 26;  A_LI_NREG_DEB  = _AB + 27
A_CH_RRSP_ROOM = _AB +  3
A_LI_RRSP_ROOM = _AB +  8
A_CH_TFSA_VAL  = _AB + 31;  A_CH_TFSA_DEB  = _AB + 32;  A_CH_TFSA_ROOM = _AB + 33
A_LI_TFSA_VAL  = _AB + 36;  A_LI_TFSA_DEB  = _AB + 37;  A_LI_TFSA_ROOM = _AB + 38

# Property group (8 cols each; base = AB + 40)
_PB = _AB + 40
A_CH_ROBERTSON_INT = _PB +  6;   A_LI_ROBERTSON_INT = _PB + 14
A_CH_MOONBROOK_INT   = _PB + 22;   A_LI_MOONBROOK_INT   = _PB + 30

# Regular group 2 (5 cols each; base = AB + 40 + 32)
_RB2 = _AB + 40 + 32
A_CH_RENT_VAL  = _RB2 + 1
A_LI_RENT_VAL  = _RB2 + 6
A_CH_OAS_VAL   = _RB2 + 11
A_LI_OAS_VAL   = _RB2 + 16
A_CH_CPP_VAL   = _RB2 + 21
A_LI_CPP_VAL   = _RB2 + 26

MIN_SIM_YEAR   = 2026   # template year (2025) is skipped

# ---------------------------------------------------------------------------
# params.h reader — keeps this script in sync without manual duplication
# ---------------------------------------------------------------------------

def _parse_params_h():
    """Return {name: numeric_value} for all #define constants in params.h."""
    defines = {}
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'params.h')
    if not os.path.exists(path):
        return defines
    with open(path, encoding='utf-8', errors='ignore') as f:
        for line in f:
            m = re.match(r'^\s*#define\s+(\w+)\s+(.*)', line)
            if not m:
                continue
            name  = m.group(1)
            value = re.sub(r'/\*.*?\*/', '', m.group(2)).strip()
            if not value:
                continue
            try:
                fv = float(eval(value))
                defines[name] = int(fv) if fv.is_integer() else fv
            except Exception:
                pass
    return defines

_params = _parse_params_h()

def _p(name, default):
    """Return a param value from params.h, falling back to *default*."""
    return _params.get(name, default)

# Simulation assumptions — sourced from params.h where possible so that
# changing params.h is the only edit required.
OAS_CLAWBACK_THRESHOLD = _p('OAS_CLAWBACK_THRESHOLD', 95323.0)  # CURRENT_YEAR (2026) base
_INFLATION              = _p('INFLATION',              0.025)
_CURRENT_YEAR           = _p('CURRENT_YEAR',           2026)

# ---------------------------------------------------------------------------
# Canadian federal + Ontario marginal rate model (pension-split optimizer)
#
# Thresholds are in approximate 2026 dollars (2024 official × 1.025²).
# They are indexed forward by _INFLATION each year when used.
# Ontario's top two thresholds ($150 K, $220 K) are legally fixed but we
# index them here for long-range consistency — the error is minor.
# Ontario surtax (~1–2 pp extra in upper brackets) is omitted; it is a
# second-order effect that rarely changes the optimal split direction.
# ---------------------------------------------------------------------------
_FED_THRESHOLDS_2026 = (60_282.0, 120_566.0, 166_537.0, 231_138.0)
_FED_RATES           = (0.150,    0.205,     0.260,     0.290,     0.330)

_ONT_THRESHOLDS_2026 = (54_069.0, 108_141.0, 157_594.0, 231_138.0)
_ONT_RATES           = (0.0505,   0.0915,    0.1116,    0.1216,    0.1316)

RRIF_CONVERSION_AGE    = 71
LIF_CONVERSION_AGE     = _p('LIF_MIX_AGE',            55)
PENSION_CREDIT_AGE     = 65
PENSION_CREDIT_MIN     = 2000.0
RALPH_RETIREMENT_YEAR  = _p('RALPH_RETIREMENT_YEAR',  2026)
SARAH_RETIREMENT_YEAR   = _p('SARAH_RETIREMENT_YEAR',   2026)
RALPH_CPP_START_AGE    = int(_p('RALPH_CPP_START_AGE',   70))
SARAH_CPP_START_AGE     = int(_p('SARAH_CPP_START_AGE',    70))
OAS_START_AGE          = int(_p('OAS_START_AGE',         70))

# Family — used for life milestone banners
SCARLET_BIRTH_YEAR     = 2009
VIOLET_BIRTH_YEAR      = 2012
# High school in Ontario: grades 9–12 (4 years).
# A child born in year Y enters grade 9 in September of year Y+14.
_HS_ENTRY_OFFSET       = 14   # age at which grade 9 begins
_HS_YEARS              = 4    # grades 9–12
_UNIV_YEARS            = 4    # typical 4-year degree

# RRIF minimum withdrawal rates by age (Ontario, matches tables.c).
# Ages 71+ use the CRA schedule; ages < 71 use the CRA formula 1/(90 - age).
_RRIF_RATES_71PLUS = {
    71:0.0528, 72:0.0540, 73:0.0553, 74:0.0567,
    75:0.0582, 76:0.0598, 77:0.0617, 78:0.0636, 79:0.0658,
    80:0.0682, 81:0.0708, 82:0.0738, 83:0.0771, 84:0.0808,
    85:0.0851, 86:0.0899, 87:0.0955, 88:0.1021, 89:0.1099,
    90:0.1192, 91:0.1306, 92:0.1449, 93:0.1634, 94:0.1879,
    95:0.2000,
}

def rrif_min_rate(age):
    """Return the CRA minimum RRIF withdrawal rate for the given age.
    Under 71: CRA formula 1/(90 - age).  71+: scheduled table rate."""
    if age < 71:
        denom = 90 - age
        return (1.0 / denom) if denom > 0 else 0.0
    for a in sorted(_RRIF_RATES_71PLUS.keys(), reverse=True):
        if age >= a:
            return _RRIF_RATES_71PLUS[a]
    return 0.0

# LIF maximum withdrawal rates by age (Ontario, matches tables.c)
_LIF_MAX_RATES = {
    55:0.065070, 56:0.065659, 57:0.066295, 58:0.066983, 59:0.067729,
    60:0.068537, 61:0.069415, 62:0.070370, 63:0.071412, 64:0.072551,
    65:0.073799, 66:0.075169, 67:0.076678, 68:0.078345, 69:0.080193,
    70:0.082250, 71:0.084548,
}

def lif_max_rate(age):
    """Return the Ontario maximum LIF withdrawal rate for the given age."""
    for a in sorted(_LIF_MAX_RATES.keys(), reverse=True):
        if age >= a:
            return _LIF_MAX_RATES[a]
    return 0.0

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
CLR_WITHDRAW_HDR = "1E5631"
CLR_INCOME_HDR   = "1565C0"
CLR_PORTF_HDR    = "4E342E"   # dark brown  — portfolio management
CLR_TAX_HDR      = "4A235A"
CLR_NOTES_HDR    = "37474F"
CLR_NW_HDR       = "1F4E79"
CLR_ROW_EVEN     = "F5F5F5"
CLR_TOTAL_WD     = "E8F5E9"
CLR_TOTAL_INC    = "E3F2FD"
CLR_NOTES_ROW    = "FFF9C4"   # pale yellow — tax note rows
CLR_NW_DATA      = "EBF3FB"
CLR_WHITE        = "FFFFFF"
CLR_TITLE_BLUE   = (0x1F, 0x4E, 0x79)
CLR_EVENT_HDR    = "7B3F00"   # dark amber  — account events header
CLR_EVENT_ROW    = "FFF3E0"   # light amber — account events row
CLR_MILE_HDR     = "4A148C"   # deep purple — milestones header
CLR_MILE_ROW     = "F3E5F5"   # light lavender — milestone rows

# ---------------------------------------------------------------------------
# Strategy-robustness thresholds  (kept in sync with retirement_to_xlsx.py)
#
# Reference: the standard retirement-planning yardstick (Bengen / Trinity /
# FIRECalc) is 95% success.  Anything below that trades safety margin for
# higher upside; below 70% the strategy fails in a meaningful fraction of
# market paths and needs revisiting.
# ---------------------------------------------------------------------------
SUCCESS_THRESHOLDS = [
    # (min_pct, label, bg_hex, fg_rgb, interpretation)
    (0.95, "STRONG",   "107C10", (0xFF, 0xFF, 0xFF),
        "At or above the 95% planning benchmark — strategy survives even "
        "unfavourable market sequences."),
    (0.85, "GOOD",     "5AA02C", (0xFF, 0xFF, 0xFF),
        "Above the 85% threshold — acceptable if you have meaningful "
        "flexibility (downsizing, part-time work, spending cuts)."),
    (0.70, "MARGINAL", "E8A33D", (0x00, 0x00, 0x00),
        "Below the 85% planning threshold — workable but sensitive to bad "
        "market sequences.  Consider later moonbrook sale, lower adjustments, "
        "or a delayed retirement date."),
    (0.50, "WEAK",     "E07020", (0xFF, 0xFF, 0xFF),
        "Under 70% survival — the strategy fails in a meaningful fraction "
        "of market paths.  Revisit assumptions before treating it as a plan."),
    (0.00, "AT RISK",  "C00000", (0xFF, 0xFF, 0xFF),
        "Under 50% — the winning strategy loses more than half the time.  "
        "Not a viable retirement plan as modelled."),
]


def classify_success_rate(rate):
    """Return (label, bg_hex, fg_rgb, interpretation) for a ratio in [0,1]."""
    for min_pct, label, bg, fg, interp in SUCCESS_THRESHOLDS:
        if rate >= min_pct:
            return label, bg, fg, interp
    return SUCCESS_THRESHOLDS[-1][1:]


def parse_success_ratio(summary_rows):
    """Pull (successes, total) from the 'Winning Strategy Successes' row."""
    for row in summary_rows:
        if not row or row[0].strip() != "Winning Strategy Successes":
            continue
        value = ""
        for cell in row[1:]:
            s = cell.strip()
            if s:
                value = s
                break
        m = re.match(r'^(\d+)\s*/\s*(\d+)$', value)
        if m:
            return int(m.group(1)), int(m.group(2))
    return None, None

# ---------------------------------------------------------------------------
# Utility helpers
# ---------------------------------------------------------------------------

def fv(row, col):
    """Safe float from a CSV column; 0.0 on any error."""
    try:
        return float(row[col])
    except (IndexError, ValueError, TypeError):
        return 0.0

def fmt_dollar(n, show_dash=True):
    if n == 0 and show_dash:
        return "\u2014"
    return f"${abs(n):,.0f}"

def fmt_pct(n):
    return f"{n * 100:.1f}%"

def rgb(hex6):
    return RGBColor(int(hex6[:2], 16), int(hex6[2:4], 16), int(hex6[4:], 16))

# ---------------------------------------------------------------------------
# XML / python-docx helpers
# ---------------------------------------------------------------------------

def add_page_number_footer(doc):
    """Add a centred 'Page X of Y' footer to the document's first section."""
    section = doc.sections[0]
    footer  = section.footer

    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _field(para, code):
        """Append a Word field (PAGE or NUMPAGES) to para."""
        run = para.add_run()
        for ftype in ('begin', 'end'):
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), ftype)
            if ftype == 'begin':
                run._r.append(fld)
                instr = OxmlElement('w:instrText')
                instr.set(qn('xml:space'), 'preserve')
                instr.text = f' {code} '
                run._r.append(instr)
            else:
                run._r.append(fld)

    p.add_run('Page ')
    _field(p, 'PAGE')
    p.add_run(' of ')
    _field(p, 'NUMPAGES')

    # Style the footer text to match the document's normal font
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def set_cell_bg(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)

def set_cell_pad(cell, top=50, bottom=50, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    mar  = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)

def set_table_borders(table, color='CCCCCC', size=4):
    tbl   = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    bds = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        bd = OxmlElement(f'w:{name}')
        bd.set(qn('w:val'),   'single')
        bd.set(qn('w:sz'),    str(size))
        bd.set(qn('w:color'), color)
        bds.append(bd)
    tblPr.append(bds)

def write_cell(cell, text, bold=False, italic=False,
               align='left', size=9, color=None):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = {
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(str(text))
    run.bold      = bold
    run.italic    = italic
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = rgb(color) if isinstance(color, str) else RGBColor(*color)

def write_hdr_cell(cell, text, bg, align='left', size=9):
    set_cell_bg(cell, bg)
    set_cell_pad(cell)
    write_cell(cell, text, bold=True, align=align, size=size, color=CLR_WHITE)

def write_data_cell(cell, text, bg=CLR_WHITE, bold=False, italic=False,
                    align='left', size=9):
    set_cell_bg(cell, bg)
    set_cell_pad(cell)
    write_cell(cell, text, bold=bold, italic=italic, align=align, size=size)

def make_table(doc, n_rows, col_widths, border_color='CCCCCC'):
    t = doc.add_table(rows=n_rows, cols=len(col_widths))
    set_table_borders(t, color=border_color)
    for i, w in enumerate(col_widths):
        for row in t.rows:
            row.cells[i].width = Inches(w)
    return t


def build_success_callout(doc, successes, total):
    """Render a prominent coloured block on the cover page showing the
    eval-pass success ratio and its threshold classification.

    Layout: a single-column, 3-row table (full page width minus margins).
      Row 1 — "Strategy Robustness" title
      Row 2 — big "X / Y paths survived  (NN%) — LABEL"
      Row 3 — italic interpretation paragraph
    """
    rate = successes / total if total > 0 else 0.0
    label, bg, fg_rgb, interp = classify_success_rate(rate)
    pct = f"{rate*100:.0f}%"

    # 7.0" = 8.5" page - 0.75" x 2 margins
    tbl = make_table(doc, n_rows=3, col_widths=[7.0], border_color=bg)

    # Title row
    c0 = tbl.cell(0, 0)
    set_cell_bg(c0, bg)
    set_cell_pad(c0, top=80, bottom=20, left=160, right=160)
    c0.text = ''
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r0 = p0.add_run("Strategy Robustness")
    r0.bold = True
    r0.font.size = Pt(11)
    r0.font.name = 'Arial'
    r0.font.color.rgb = RGBColor(*fg_rgb)

    # Big number row
    c1 = tbl.cell(1, 0)
    set_cell_bg(c1, bg)
    set_cell_pad(c1, top=20, bottom=20, left=160, right=160)
    c1.text = ''
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run(
        f"{successes} / {total} market paths survived  ({pct})   \u2014   {label}"
    )
    r1.bold = True
    r1.font.size = Pt(18)
    r1.font.name = 'Arial'
    r1.font.color.rgb = RGBColor(*fg_rgb)

    # Interpretation row
    c2 = tbl.cell(2, 0)
    set_cell_bg(c2, bg)
    set_cell_pad(c2, top=20, bottom=80, left=160, right=160)
    c2.text = ''
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run(interp)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.name = 'Arial'
    r2.font.color.rgb = RGBColor(*fg_rgb)

    # Small spacer paragraph under the callout so the stochastic-returns note
    # doesn't butt right up against it.
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)

def add_sub_label(doc, text, r, g, b, space_before=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(8)
    run.font.name      = 'Arial'
    run.font.color.rgb = RGBColor(r, g, b)

# ---------------------------------------------------------------------------
# Table widths (inches): label col + Ralph + Sarah + Combined
# ---------------------------------------------------------------------------
W4 = [3.0, 1.2, 1.2, 1.1]   # 4-column tables
W3 = [2.9, 1.55, 1.55]       # 3-column tables (tax)
W2 = [4.2, 2.3]              # 2-column tables (notes)

# ---------------------------------------------------------------------------
# Milestone helpers
# ---------------------------------------------------------------------------

def compute_all_milestones(moonbrook_sale_year, robertson_payoff_year, moonbrook_payoff_year,
                           survivor_name=None, survivor_year=None, survivor_age=None):
    """
    Build a dict mapping calendar year → list of milestone strings.
    Academic years (Sept–June) are attributed to the September-start year
    for "starts" events, and to the June-end year for "graduates" events.

    survivor_name / survivor_year / survivor_age: when a --survivor run was used,
    inject the death event into the milestone banner for that year.
    """
    ms = {}

    def add(year, text):
        ms.setdefault(year, []).append(text)

    # ---- Survivor / death event ----
    if survivor_name and survivor_year:
        surviving = 'Sarah' if survivor_name == 'Ralph' else 'Ralph'
        age_str   = f' (age {survivor_age})' if survivor_age else ''
        add(survivor_year,
            f'\U0001f397  {survivor_name} passes away{age_str} \u2014 {surviving} continues as survivor')

    # ---- Retirement ----
    if RALPH_RETIREMENT_YEAR == SARAH_RETIREMENT_YEAR:
        add(RALPH_RETIREMENT_YEAR, '\U0001f3e6  Ralph & Sarah retire')
    else:
        add(RALPH_RETIREMENT_YEAR, '\U0001f3e6  Ralph retires')
        add(SARAH_RETIREMENT_YEAR,  '\U0001f3e6  Sarah retires')

    # ---- Children ----
    for name, birth in (('Scarlet', SCARLET_BIRTH_YEAR),
                        ('Violet',  VIOLET_BIRTH_YEAR)):
        hs_start  = birth + _HS_ENTRY_OFFSET          # Sept of grade-9 year
        hs_end    = hs_start + _HS_YEARS               # June of grade-12 year
        univ_end  = hs_end   + _UNIV_YEARS

        # Only emit HS-start if it falls within the simulated window
        if hs_start >= MIN_SIM_YEAR:
            add(hs_start, f'\U0001f3eb  {name} starts high school (Grade 9)')

        grade12_start = hs_end - 1   # September of Grade-12 year
        if grade12_start >= MIN_SIM_YEAR:
            add(grade12_start, f'\U0001f4da  {name} enters Grade 12 — final year of high school')

        add(hs_end,  f'\U0001f393  {name} graduates high school')
        add(hs_end,  f'\U0001f3eb  {name} starts university — Year 1 of {_UNIV_YEARS}')
        for yr_n in range(2, _UNIV_YEARS + 1):
            add(hs_end + yr_n - 1,
                f'\U0001f4d6  {name} — University Year {yr_n} of {_UNIV_YEARS}')
        add(univ_end, f'\U0001f393  {name} graduates university')

    # ---- Mortgages ----
    if moonbrook_sale_year:
        add(moonbrook_sale_year, '\U0001f3e0  Moonbrook Street property sold')
    elif moonbrook_payoff_year:
        add(moonbrook_payoff_year, '\U0001f3e0  Moonbrook Street mortgage fully paid off')

    if robertson_payoff_year:
        add(robertson_payoff_year, '\U0001f3e0  Robertson mortgage fully paid off')

    return ms


def build_milestones_banner(doc, milestones):
    """
    Render a compact purple milestone banner for the given year.
    milestones is a list of strings; does nothing if the list is empty.
    """
    if not milestones:
        return

    add_sub_label(doc, 'Milestones', 0x4A, 0x14, 0x8C, space_before=0)

    t = make_table(doc, 1 + len(milestones), [6.5], border_color='9C27B0')

    hdr = t.rows[0].cells
    write_hdr_cell(hdr[0], 'Life Events This Year', CLR_MILE_HDR)

    for i, text in enumerate(milestones):
        cell = t.rows[i + 1].cells[0]
        set_cell_bg(cell, CLR_MILE_ROW)
        set_cell_pad(cell)
        write_cell(cell, text, size=9)


# ---------------------------------------------------------------------------
# Front-matter pages  (aging-friendly guidance sections)
#
# These are static reference pages inserted between the cover page and the
# year-by-year action plan.  They are written in plain language and sized
# large enough to remain legible later in life, when complex multi-column
# tables become harder to parse.
# ---------------------------------------------------------------------------

def _section_heading(doc, text, color=CLR_TITLE_BLUE):
    """Emit a consistent section title used across the front-matter pages."""
    h = doc.add_heading(text, level=1)
    h.runs[0].font.size      = Pt(16)
    h.runs[0].font.color.rgb = RGBColor(*color)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after  = Pt(6)
    return h


def _intro_para(doc, text, size=10, italic=False, color=(0x40, 0x40, 0x40)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.size      = Pt(size)
    r.font.name      = 'Arial'
    r.italic         = italic
    r.font.color.rgb = RGBColor(*color)
    return p


def _bullet(doc, text, indent=0.25, size=10, bold_lead=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(indent)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(3)
    if bold_lead and '—' in text:
        lead, rest = text.split('—', 1)
        r1 = p.add_run('\u2022  ' + lead.strip() + '  — ')
        r1.bold = True; r1.font.size = Pt(size); r1.font.name = 'Arial'
        r2 = p.add_run(rest.strip())
        r2.font.size = Pt(size); r2.font.name = 'Arial'
    else:
        r = p.add_run('\u2022  ' + text)
        r.font.size = Pt(size); r.font.name = 'Arial'
    return p


def build_table_of_contents(doc):
    """Insert a Word TOC field so the reader can jump to any year."""
    _section_heading(doc, 'Table of Contents')
    _intro_para(doc,
        'This plan is organised into front-matter reference pages followed '
        'by a dedicated page for each calendar year of the simulation.  '
        'To refresh the page numbers shown below: right-click anywhere in '
        'the table, choose Update Field, then select "Update entire table."',
        italic=True, size=9, color=(0x55, 0x55, 0x55))

    # Insert a Word TOC field targeting Heading 1 (years + section titles)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \o "1-1" \h \z \u '
    run._r.append(instr)

    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)

    placeholder = OxmlElement('w:t')
    placeholder.text = 'Right-click and choose "Update Field" to populate.'
    run._r.append(placeholder)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)

    doc.add_page_break()


def build_rules_page(doc):
    """Plain-language rules to follow year after year."""
    _section_heading(doc, 'Rules to Follow')
    _intro_para(doc,
        'These are the core rules that should guide every year\u2019s decisions.  '
        'When in doubt, fall back to these — they survive almost every market '
        'scenario the simulator explored.')

    rules = [
        'Follow the plan for the current year first — do not improvise until you have read the whole page.',
        'Never take on new debt in retirement.  Pay credit cards in full every month.',
        'Keep at least 12 months of expenses in the chequing/savings buffer before investing surplus cash.',
        'Withdraw from accounts in this priority order: Non-Registered → RRIF/LIF (minimums) → TFSA (last).  This preserves tax-free growth.',
        'In a down market year: pull the RRIF before the TFSA to shrink next year\u2019s mandatory minimum and preserve tax-free growth.',
        'Land taxable income just below the OAS clawback threshold wherever possible.  The 15\u00a2/$ OAS recovery tax is avoidable with planning.',
        'Sign the T1032 pension-split election every April.  It almost always saves combined tax.',
        'Have a licenced tax professional review your return before filing, every single year.  Fees are far smaller than one missed credit.',
        'Update beneficiary designations on every account after any major life event (death, divorce, new grandchild).',
        'Re-read the "Red Flags" page below any time something feels wrong.  When in doubt, call before acting.',
    ]
    for r in rules:
        _bullet(doc, r)

    doc.add_page_break()


def build_red_flags_page(doc):
    """Warning signs that should trigger a phone call."""
    _section_heading(doc, 'Red Flags \u2014 Stop and Ask for Help', color=(0xC0, 0x00, 0x00))
    _intro_para(doc,
        'If you notice ANY of the following, pause before doing anything and '
        'call your financial advisor or a trusted family member.  '
        'Scams and costly mistakes almost always show up as one of these.',
        color=(0x70, 0x00, 0x00))

    flags = [
        'Someone phones, texts, or emails saying your CRA account is at risk, or demanding taxes be paid with gift cards, wire transfers, or cryptocurrency.  This is always a scam.  CRA never calls to demand immediate payment.',
        'A "new advisor" or stranger wants you to move money into an investment with guaranteed returns of 8% or more.  Real investments never come with a guarantee.',
        'You are asked to sign a document you have not read, or the person pushing you says "sign now before the offer expires."  Legitimate deals do not expire overnight.',
        'Your bank or RRIF balance drops by more than 10% in a single statement that you do not recognise.  Log in, check every transaction, and call the bank fraud line if anything is unfamiliar.',
        'You cannot remember if you already paid a bill or made a withdrawal this month.  Stop, open the banking app, and check before acting.',
        'Anything in this plan no longer seems to describe reality (e.g., account balances look wildly different from what is shown).  Regenerate the plan or ask your advisor to walk through the changes.',
        'You feel rushed, confused, or pressured — by anyone, for any reason.  The correct response is always "I need to think about this; I will call you back tomorrow."',
    ]
    for f in flags:
        _bullet(doc, f)

    _intro_para(doc,
        'If the situation involves money leaving an account and you are unsure, '
        'call the bank\u2019s fraud line FIRST (the number on the back of your card) '
        'before anything else.  It is always free and usually resolves in minutes.',
        italic=True, size=9, color=(0x70, 0x00, 0x00))

    doc.add_page_break()


def build_contacts_page(doc):
    """Blank template for key phone numbers — user fills in the detail."""
    _section_heading(doc, 'Who to Call')
    _intro_para(doc,
        'Fill this page in once, keep the printed copy near the phone, and '
        'update it whenever any of these contacts change.  These are the '
        'only numbers Ralph or Sarah should need in an emergency.')

    contacts = [
        ('Financial Advisor',          'Primary point of contact for all investment decisions and tax questions.'),
        ('Accountant / Tax Preparer',  'Files the annual T1 return; owns the T1032 pension-split paperwork.'),
        ('Family Doctor',              ''),
        ('Lawyer (wills, power of attorney)', ''),
        ('Executor of the Will',       ''),
        ('Power of Attorney \u2014 Finance', 'Authorised to act on banking/investing if either of you cannot.'),
        ('Power of Attorney \u2014 Health',  'Authorised to make medical decisions if either of you cannot.'),
        ('Primary Bank \u2014 Fraud Line',   'On the back of each debit / credit card.'),
        ('Bank \u2014 Branch Manager',       ''),
        ('Service Canada (CPP / OAS)',       '1-800-277-9914'),
        ('CRA Individual Enquiries',         '1-800-959-8281'),
        ('Scarlet',                    ''),
        ('Violet',                     ''),
        ('Trusted Neighbour / Friend', 'Physically nearby; can check in and drive to appointments.'),
    ]

    t = make_table(doc, 1 + len(contacts), [2.3, 2.0, 2.7])

    hdr = t.rows[0].cells
    write_hdr_cell(hdr[0], 'Role',       CLR_NW_HDR)
    write_hdr_cell(hdr[1], 'Name / Phone', CLR_NW_HDR)
    write_hdr_cell(hdr[2], 'Notes',      CLR_NW_HDR)

    for i, (role, note) in enumerate(contacts):
        bg = CLR_ROW_EVEN if i % 2 == 0 else CLR_WHITE
        cells = t.rows[i + 1].cells
        write_data_cell(cells[0], role, bg=bg, bold=True, size=9)
        write_data_cell(cells[1], '',   bg=bg, size=9)     # user fills in
        write_data_cell(cells[2], note, bg=bg, italic=True, size=8)

    doc.add_page_break()


def build_glossary_page(doc):
    """Plain-language definitions of every acronym used in the plan."""
    _section_heading(doc, 'Glossary')
    _intro_para(doc,
        'Every abbreviation used in the year-by-year pages is defined here.  '
        'If a page refers to a term you do not recognise, flip back to this '
        'glossary before acting on anything.')

    terms = [
        ('RRSP',      'Registered Retirement Savings Plan.  Pre-tax retirement account; '
                      'contributions reduce this year\u2019s taxable income, withdrawals are taxed.'),
        ('RRIF',      'Registered Retirement Income Fund.  The RRSP must be converted to a RRIF '
                      'by December 31 of the year you turn 71.  A minimum percentage must be '
                      'withdrawn every year thereafter.'),
        ('DCPP',      'Defined Contribution Pension Plan.  Employer pension that behaves like '
                      'an RRSP — balance grows; withdrawals are taxed as income.'),
        ('LIRA',      'Locked-In Retirement Account.  Essentially a locked RRSP created when '
                      'you leave an employer with a pension.  Must be converted to a LIF.'),
        ('LIF',       'Life Income Fund.  The retirement-phase version of a LIRA.  Has a minimum '
                      'AND a maximum withdrawal each year (unlike a RRIF, which only has a minimum).'),
        ('TFSA',      'Tax-Free Savings Account.  No tax on growth or withdrawals.  '
                      'Withdrawn amounts are re-added to your contribution room on January 1 next year.'),
        ('CPP',       'Canada Pension Plan.  Monthly government pension based on contributions made '
                      'during working years.  Can start anywhere from age 60 to 70.'),
        ('OAS',       'Old Age Security.  Monthly government pension funded by general tax revenue.  '
                      'Starts at 65 unless deferred; subject to a 15% clawback on income above a threshold.'),
        ('BPA',       'Basic Personal Amount.  The first ~$16K of income (indexed) is federally '
                      'tax-free for every Canadian.  Ontario has its own similar amount.'),
        ('ACB',       'Adjusted Cost Base.  The tax "purchase price" of a non-registered investment.  '
                      'Capital gain = sale proceeds \u2212 ACB.  Resetting the ACB periodically '
                      '("harvesting") can reduce future taxes.'),
        ('CCA',       'Capital Cost Allowance.  Annual depreciation deduction on a rental property.'),
        ('UCC',       'Undepreciated Capital Cost.  The running balance of the property\u2019s CCA pool; '
                      'shortfall at sale time is taxed as recapture (full income, not capital gain).'),
        ('T1032',     'Pension Income Splitting election form.  Lets a higher-income spouse transfer '
                      'up to 50% of eligible pension income to the lower-income spouse to equalise tax.'),
        ('Schedule 3','The T1 schedule where capital gains and losses are reported.  Attach transaction '
                      'details: date, proceeds, adjusted cost base.'),
        ('T776 / T2125', 'Rental income (T776) and self-employment (T2125) tax forms.  '
                      'Used when Moonbrook Street rental income or CCA recapture must be reported.'),
        ('Clawback',  'OAS Recovery Tax.  When taxable income exceeds an indexed threshold '
                      '($95,323 in 2026, inflation-indexed annually), OAS is reduced by 15\u00a2 for every dollar above it.'),
        ('Installments', 'Quarterly prepayments of income tax that CRA demands whenever tax owing '
                      'exceeds $3,000 in the current year AND in either of the two preceding years.'),
    ]

    t = make_table(doc, 1 + len(terms), [1.3, 5.7])
    hdr = t.rows[0].cells
    write_hdr_cell(hdr[0], 'Term',       CLR_NW_HDR)
    write_hdr_cell(hdr[1], 'Definition', CLR_NW_HDR)

    for i, (term, defn) in enumerate(terms):
        bg = CLR_ROW_EVEN if i % 2 == 0 else CLR_WHITE
        cells = t.rows[i + 1].cells
        write_data_cell(cells[0], term, bg=bg, bold=True, size=9)
        write_data_cell(cells[1], defn, bg=bg, size=9)

    doc.add_page_break()


def build_standard_year_checklist(doc):
    """Template checklist to run through every calendar year."""
    _section_heading(doc, 'Standard Year Checklist')
    _intro_para(doc,
        'Run through this checklist every January, April, and December.  '
        'The year-by-year pages that follow customise this list with the '
        'exact dollar amounts and forms for each year — but these are the '
        'items that repeat every single year.')

    sections = [
        ('January \u2014 Start of the Year', [
            'Review the plan page for the upcoming calendar year (age, expected income, withdrawals).',
            'Confirm automatic RRIF and LIF payments are scheduled — and match the planned amounts.',
            'Transfer last-year\u2019s TFSA withdrawals back into TFSA room (room resets Jan 1).',
            'Record the new Federal & Ontario tax thresholds published by CRA.',
        ]),
        ('April \u2014 Tax Filing Month', [
            'Gather all T4A (pension), T4RIF (RRIF), T4OAS, T4AP (CPP), T5 (investment), T3 (trust) slips.',
            'Sign Form T1032 for pension-income splitting — both spouses.',
            'Confirm Schedule 3 (capital gains) includes every non-reg disposition during the year.',
            'File by April 30.  Set up CRA direct deposit for any refund.',
            'If CRA is sending installment reminders: a voluntary lump-sum payment does NOT exit the regime (it reduces your April balance, not your net tax owing).  Instead, elect the "current-year option" if this year\'s tax will be lower than last year\'s, or increase source withholding on RRIF / CPP / OAS via Form T1213(OAS) or ISP-3520 — withholding counts against net tax owing, installments do not.',
        ]),
        ('Mid-Year \u2014 June / July', [
            'Review the portfolio with your financial advisor.  Rebalance only if allocations have drifted 5% or more.',
            'Confirm the year-to-date spending matches the planned Annual Expenses for this year; adjust if off by more than 10%.',
        ]),
        ('December \u2014 Year-End Tax Planning', [
            'Size the year-end RRIF top-up so total taxable income lands just below the OAS clawback threshold (or the Age Amount phaseout threshold, whichever is binding).',
            'Harvest non-registered capital gains if current-year income leaves room below the top-bracket threshold (50% inclusion rate applies below ~$250K).',
            'Make TFSA contributions if any room is unused (room is published on CRA My Account).',
            'Book the annual tax-review meeting with your accountant for March.',
        ]),
    ]

    for title, items in sections:
        sh = doc.add_paragraph()
        sh.paragraph_format.space_before = Pt(6)
        sh.paragraph_format.space_after  = Pt(3)
        sr = sh.add_run(title)
        sr.bold = True
        sr.font.size = Pt(11); sr.font.name = 'Arial'
        sr.font.color.rgb = RGBColor(*CLR_TITLE_BLUE)
        for it in items:
            _bullet(doc, it, indent=0.25, size=10)

    doc.add_page_break()


def build_surviving_spouse_page(doc):
    """Contingency instructions for whichever spouse survives the other."""
    _section_heading(doc, 'If One Spouse Passes Away', color=(0x4A, 0x14, 0x8C))
    _intro_para(doc,
        'This page describes the steps the surviving spouse should take, '
        'regardless of who passes away first.  Keep this page bookmarked and '
        'printed.  The simulator can also regenerate this plan on a single-'
        'survivor basis; ask the advisor for the survivor variant if needed.')

    immediate = [
        'Call the financial advisor and the accountant within the first two weeks.  They will coordinate the paperwork.',
        'Obtain 10\u201315 certified copies of the death certificate from the funeral home.  Every institution needs one.',
        'Notify Service Canada: CPP survivor benefit and OAS; they cancel the deceased\u2019s benefits and start the survivor benefit.',
        'Notify every bank and investment institution.  They will freeze joint accounts briefly and re-title them to the survivor.',
        'Notify the life-insurance carrier(s) to begin the claim.',
        'Do NOT close any account or sell any asset in the first 30 days unless the advisor specifically approves it.',
    ]

    financial = [
        'RRSP / RRIF: roll over tax-free to the surviving spouse if named as the beneficiary or successor annuitant.  Named beneficiary avoids probate.',
        'TFSA: the survivor can roll the deceased\u2019s TFSA into their own TFSA without using personal contribution room, if named as successor holder.',
        'Non-registered investments: deemed sold at fair market value at date of death; capital gains may be deferred by spousal rollover.',
        'Principal residence (Robertson): passes to survivor tax-free if jointly held.',
        'Moonbrook Street rental: triggers deemed disposition and possible CCA recapture on the deceased\u2019s half.  This is the largest tax exposure and must be modelled by the accountant.',
        'CPP: survivor benefit is available; amount depends on both spouses\u2019 contribution records.  Apply within 12 months.',
        'OAS: survivor benefit is NOT automatic for spouses 65 and over — each gets their own OAS already.  There is a small widowed allowance for 60\u201364 in some cases.',
    ]

    ongoing = [
        'Tax brackets change dramatically: the survivor now files a single return at single-bracket rates.  Pension splitting (T1032) is no longer available.  This may push more of the income into higher brackets.',
        'Run the plan with --survivor to see an updated year-by-year plan tailored to the single-spouse scenario.',
        'Revisit the will within 6 months.  Update beneficiaries on every RRIF, TFSA, and insurance policy.',
        'Consider whether downsizing (Robertson sale) should be accelerated; the senior-living transition may make more sense sooner without a partner.',
    ]

    for title, items in [
        ('Immediate Steps (Within 30 Days)', immediate),
        ('Financial Transfers', financial),
        ('Ongoing Changes to This Plan', ongoing),
    ]:
        sh = doc.add_paragraph()
        sh.paragraph_format.space_before = Pt(6)
        sh.paragraph_format.space_after  = Pt(3)
        sr = sh.add_run(title)
        sr.bold = True; sr.font.size = Pt(11); sr.font.name = 'Arial'
        sr.font.color.rgb = RGBColor(0x4A, 0x14, 0x8C)
        for it in items:
            _bullet(doc, it, indent=0.25, size=10)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Life-transition callouts
#
# Certain years involve one-time decisions that deserve a distinct, loudly
# coloured panel on the year page — distinct from the smaller milestone
# banner (which is just a reminder) and the account-change table (which is
# scoped to the mechanics of opening a RRIF or LIF).  These callouts give
# context and decision-support for the full transition.
# ---------------------------------------------------------------------------

def build_transition_callout(doc, year, ch_age, li_age,
                             moonbrook_sale_year, robertson_payoff_year):
    """
    Emit a large amber/blue callout describing the key decision for the year
    when the simulated year is one of a small set of life-transition years.
    Returns True when a callout was emitted (so the caller can skip any
    redundant milestone-banner rendering, if desired).
    """
    RALPH_OAS_YEAR = 1961 + OAS_START_AGE  # derived from params.h OAS_START_AGE
    # Use a simple year table driven by ages: any year where Ralph or Sarah
    # cross one of the well-known thresholds.
    events = []

    # ---- 2026 / retirement entry year ----
    if year == RALPH_RETIREMENT_YEAR or year == SARAH_RETIREMENT_YEAR:
        events.append((
            'Entering Retirement \u2014 Foundations',
            [
                'Transfer each DCPP balance to a Locked-In Retirement Account (LIRA) at your financial institution.  Ontario law requires you to be at least 55 before a LIRA can be converted to a Life Income Fund (LIF); until then the funds remain locked.  At age 55, up to 50% of the LIRA balance can be unlocked into an RRSP — the simulation applies this split in the age-55 year.',
                'Set up a dedicated retirement chequing account; all withdrawals will flow through it.',
                'Subscribe to CRA My Account and Service Canada online; all future correspondence will be digital.',
                'Book an annual review with your financial advisor (first Friday of February is a common standing date).',
                'Review life-insurance and disability-insurance policies: cancel any that were tied to employment.',
            ]
        ))

    # ---- Moonbrook Street sale year ----
    if moonbrook_sale_year and year == moonbrook_sale_year:
        events.append((
            'Moonbrook Street Sale Year \u2014 The Biggest Tax Event',
            [
                'Engage a real-estate lawyer at least 60 days before the planned closing.',
                'Capital gain = sale proceeds \u2212 adjusted cost base.  50% of the gain up to the first $250K per person, then 66.67% above.',
                'CCA recapture: every dollar of depreciation previously deducted comes back as FULL-RATE income.  This is usually the largest hidden tax.',
                'Bracket-smoothing lever: closing in late December vs. early January splits the capital-gain income across two tax years, potentially halving the marginal rate on the upper portion.  Discuss this timing with your lawyer and accountant before finalising the closing date.',
                'Consider making large RRSP contributions this year (using any carry-forward room) to offset the income spike.',
                'Use Form T776 for the final year\u2019s rental income and to declare the disposition.',
                'Do NOT withdraw from RRSP/RRIF above the minimum this year — capital-gain income will already push you into the top brackets.',
            ]
        ))

    # ---- Ralph CPP/OAS start (age 65) ----
    if ch_age == 65:
        events.append((
            'Ralph Turns 65 \u2014 CPP, OAS, and Pension Credit',
            [
                (f'OAS is deferred to age {OAS_START_AGE} in this plan \u2014 each month deferred past 65 adds '
                 f'0.6\u00a0%/month; deferring from 65 to {OAS_START_AGE} increases the benefit by '
                 f'{(OAS_START_AGE - 65) * 12 * 0.6:.0f}\u00a0%.  '
                 f'Apply at Service Canada 11 months before your {OAS_START_AGE}th birthday.'
                 if OAS_START_AGE > 65 else
                 'Apply for OAS at Service Canada now \u2014 OAS starts this year.'),
                (f'CPP is deferred to age {RALPH_CPP_START_AGE} in this plan \u2014 no application needed yet.  '
                 f'Each month deferred past 65 adds 0.7\u00a0%/month to the benefit; '
                 f'deferring to {RALPH_CPP_START_AGE} increases it by {(RALPH_CPP_START_AGE - 65) * 12 * 0.7:.0f}%.'
                 if RALPH_CPP_START_AGE > 65 else
                 'Apply for CPP \u2014 it starts this year in the plan.'),
                'The $2,000 Pension Income Credit becomes claimable on the first $2,000 of eligible pension income (RRIF/LIF).',
                'T1032 pension-income splitting becomes available now for Ralph\u2019s RRIF/LIF withdrawals.',
                'Age Amount credit (federal + Ontario) begins to phase in; it phases out above ~$44K income (2026 base, inflation-indexed annually).',
            ]
        ))

    # ---- Sarah CPP/OAS start (age 65) ----
    if li_age == 65:
        events.append((
            'Sarah Turns 65 \u2014 CPP, OAS, and Pension Credit',
            [
                (f'OAS is deferred to age {OAS_START_AGE} in this plan \u2014 each month deferred past 65 adds '
                 f'0.6\u00a0%/month; deferring from 65 to {OAS_START_AGE} increases the benefit by '
                 f'{(OAS_START_AGE - 65) * 12 * 0.6:.0f}\u00a0%.  '
                 f'Apply at Service Canada 11 months before your {OAS_START_AGE}th birthday.'
                 if OAS_START_AGE > 65 else
                 'Apply for OAS at Service Canada now \u2014 OAS starts this year.'),
                (f'CPP is deferred to age {SARAH_CPP_START_AGE} in this plan \u2014 no application needed yet.  '
                 f'Each month deferred past 65 adds 0.7\u00a0%/month to the benefit; '
                 f'deferring to {SARAH_CPP_START_AGE} increases it by {(SARAH_CPP_START_AGE - 65) * 12 * 0.7:.0f}%.'
                 if SARAH_CPP_START_AGE > 65 else
                 'Apply for CPP \u2014 it starts this year in the plan.'),
                'The $2,000 Pension Income Credit becomes claimable on the first $2,000 of eligible pension income.',
                'Sarah\u2019s RRIF/LIF withdrawals now qualify for T1032 pension-income splitting.',
            ]
        ))

    # ---- OAS application reminder (one year before OAS_START_AGE) ----
    for _name, _age in (('Ralph', ch_age), ('Sarah', li_age)):
        if _age == OAS_START_AGE - 1:
            _defer_pct = (OAS_START_AGE - 65) * 12 * 0.6
            events.append((
                f'{_name} Turns {OAS_START_AGE - 1} \u2014 Apply for OAS Now',
                [
                    f'Apply for OAS at Service Canada this year \u2014 payments begin at age {OAS_START_AGE}.',
                    (f'Deferring from 65 to {OAS_START_AGE} increased the monthly benefit by {_defer_pct:.0f}\u00a0%.  '
                     f'Apply online at canada.ca/en/services/benefits/publicpensions or call 1\u2011800\u2011277\u20119914.'
                     if OAS_START_AGE > 65 else
                     'Apply online at canada.ca/en/services/benefits/publicpensions or call 1\u2011800\u2011277\u20119914.'),
                    'Have your SIN, banking direct-deposit details, and date-of-birth documentation ready.',
                    'Confirm the first expected payment date and set a calendar reminder to follow up if it does not arrive.',
                ]
            ))

    # ---- RRIF mandatory conversion (age 71 in the year following) ----
    # Handled by existing build_account_events, but add a big-picture note
    # if the age is exactly 70 (one year before mandatory conversion)
    for name, age in (('Ralph', ch_age), ('Sarah', li_age)):
        if age == 70:
            events.append((
                f'{name} Turns 70 \u2014 RRIF Conversion Approaches',
                [
                    f'{name}\u2019s RRSP MUST be fully converted to a RRIF (or annuity) before December 31 of next year.',
                    'Contact your financial institution now to schedule the conversion — avoid the last-week December rush.',
                    'If the other spouse is younger, elect to use THEIR age for the RRIF minimum percentage to defer tax.  This election can only be made at conversion.',
                    'No tax is triggered by the conversion itself; it is purely an administrative change.',
                ]
            ))

    # ---- Robertson payoff year ----
    if robertson_payoff_year and year == robertson_payoff_year:
        events.append((
            'Robertson Mortgage Paid Off',
            [
                'Cancel the automatic mortgage payment at your bank.',
                'Redirect the freed cash flow to the non-registered portfolio or an extra TFSA contribution.',
                'Keep the home-insurance policy and property-tax pre-authorised debit running.',
                'Request a discharge statement from the lender and file it with the property papers.',
            ]
        ))

    if not events:
        return False

    add_sub_label(doc, 'Life Transition This Year \u2014 Action Required',
                  0x1F, 0x4E, 0x79, space_before=0)

    for title, bullets in events:
        # One table per event: a single-cell heading row, then a bullet cell
        t = make_table(doc, 2, [6.5], border_color='1F4E79')
        hdr = t.rows[0].cells[0]
        set_cell_bg(hdr, '1F4E79'); set_cell_pad(hdr, top=60, bottom=60)
        write_cell(hdr, title, bold=True, size=10, color=CLR_WHITE)

        body = t.rows[1].cells[0]
        set_cell_bg(body, 'E8F1FA'); set_cell_pad(body, top=80, bottom=80, left=140, right=140)
        body.text = ''
        first = True
        for b in bullets:
            if first:
                p = body.paragraphs[0]; first = False
            else:
                p = body.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run('\u2022  ' + b)
            r.font.size = Pt(9); r.font.name = 'Arial'

    return True


# ---------------------------------------------------------------------------
# 0. Account Events  (one-time openings and conversions)
# ---------------------------------------------------------------------------

def build_account_events(doc, row,
                         prev_ch_rrsp=0.0, prev_li_rrsp=0.0,
                         ch_rrif_open=False, li_rrif_open=False,
                         ch_lif_open=False, li_lif_open=False):
    """
    Emits a highlighted callout for one-time account transitions:
      • DCPP / LIRA  →  LIF (first year DCPP withdrawal appears)
      • RRSP         →  RRIF on first withdrawal year (lower fees; voluntary early)
      • RRSP         →  RRIF at age 71  (mandatory conversion deadline)
    Returns (emitted, ch_rrif_open, li_rrif_open, ch_lif_open, li_lif_open).
    """
    ch_age = int(fv(row, C_CH_AGE))
    li_age = int(fv(row, C_LI_AGE))

    events = []   # list of (title, detail)

    # ---- DCPP / LIRA  →  LIF ----
    # Trigger: first year the DCPP debit turns positive (covers both the normal
    # case where retirement precedes LIF age and the delayed case where the
    # person retires *after* LIF age and the transfer fires at retirement).
    for name, age, dcpp_deb_col, dcpp_val_col, already_open in [
        ('Ralph', ch_age, A_CH_DCPP_DEB, A_CH_DCPP_VAL, ch_lif_open),
        ('Sarah',  li_age, A_LI_DCPP_DEB, A_LI_DCPP_VAL, li_lif_open),
    ]:
        dcpp_deb = fv(row, dcpp_deb_col)
        if dcpp_deb <= 0 or already_open:
            continue
        # The simulation halves the DCPP balance before the first withdrawal;
        # the CSV value is already the post-split LIF balance.
        lif_bal   = fv(row, dcpp_val_col)
        transfer  = lif_bal          # both halves are equal
        original  = lif_bal * 2      # pre-split total
        max_pct   = lif_max_rate(age)
        max_amt   = lif_bal * max_pct
        events.append((
            f'{name} \u2014 Open LIF Account (DCPP / LIRA Conversion)',
            f'{name}\'s DCPP / LIRA balance of {fmt_dollar(original)} '
            f'has been split 50 / 50 this year:\n'
            f'  \u2022  {fmt_dollar(transfer)} transferred to {name}\'s RRSP '
            f'(now unlocked, flexible withdrawal).\n'
            f'  \u2022  {fmt_dollar(lif_bal)} remains as a Life Income Fund (LIF).\n'
            f'Contact your plan administrator to formally open the LIF account.\n'
            f'LIF max withdrawal this year: {fmt_pct(max_pct)} = {fmt_dollar(max_amt)}  '
            f'(Ontario schedule).  There is also a CRA RRIF minimum floor.'
        ))
        if name == 'Ralph': ch_lif_open = True
        else:               li_lif_open = True

    # ---- RRSP  →  RRIF: first withdrawal year (early, voluntary, lower fees) ----
    ch_rrsp_now = fv(row, C_CH_RRSP)
    li_rrsp_now = fv(row, C_LI_RRSP)

    for name, age, rrsp_val_col, rrsp_now, rrsp_prev, already_open, use_spouse_age, spouse_age in [
        ('Ralph', ch_age, A_CH_RRSP_VAL, ch_rrsp_now, prev_ch_rrsp,
         ch_rrif_open, False, li_age),
        ('Sarah',  li_age, A_LI_RRSP_VAL, li_rrsp_now, prev_li_rrsp,
         li_rrif_open, True,  ch_age),
    ]:
        first_withdrawal = (rrsp_now > 0 and rrsp_prev <= 0)
        if not first_withdrawal or already_open or age >= RRIF_CONVERSION_AGE:
            continue
        rrsp_bal = fv(row, rrsp_val_col)
        rate_age = spouse_age if use_spouse_age else age
        min_pct  = rrif_min_rate(rate_age)
        min_amt  = rrsp_bal * min_pct
        spouse_note = (
            f'  \u2022  Sarah should elect to use Ralph\'s age ({spouse_age}) for '
            f'RRIF minimums — rate {fmt_pct(rrif_min_rate(age))} (own age) vs '
            f'{fmt_pct(min_pct)} (Ralph\'s age).  This election must be '
            f'declared when the RRIF is opened and cannot be changed later.\n'
        ) if use_spouse_age else ''
        events.append((
            f'{name} \u2014 Open RRIF Now (Early, Voluntary)',
            f'{name} is already making RRSP withdrawals this year.  '
            f'Opening a RRIF now avoids RRSP redemption fees and gives more flexibility:\n'
            f'  \u2022  Transfer {name}\'s RRSP balance ({fmt_dollar(rrsp_bal)}) '
            f'into a new RRIF at your financial institution.\n'
            f'  \u2022  No tax is triggered by the conversion itself.\n'
            f'{spouse_note}'
            f'  \u2022  RRIF minimum this year: {fmt_pct(min_pct)} = {fmt_dollar(min_amt)} '
            f'— well below the planned withdrawal, so no extra income is forced.\n'
            f'  \u2022  RRSP must be fully converted to RRIF by age 71 regardless.'
        ))
        # Update open flags
        if name == 'Ralph': ch_rrif_open = True
        else:               li_rrif_open = True

    # ---- RRSP  →  RRIF  (mandatory at age RRIF_CONVERSION_AGE) ----
    for name, age, rrsp_col, use_spouse_age, spouse_age, already_open in [
        ('Ralph', ch_age, A_CH_RRSP_VAL, False, li_age, ch_rrif_open),
        ('Sarah',  li_age, A_LI_RRSP_VAL, True,  ch_age, li_rrif_open),
    ]:
        if age != RRIF_CONVERSION_AGE:
            continue
        rrsp_bal  = fv(row, rrsp_col)
        rate_age  = spouse_age if use_spouse_age else age
        min_pct   = rrif_min_rate(rate_age)
        min_amt   = rrsp_bal * min_pct
        if already_open:
            # RRIF was opened early; just note the mandatory minimum rate change
            events.append((
                f'{name} \u2014 RRIF Mandatory Minimum Rate Change (Age 71)',
                f'{name}\'s RRIF minimum rate switches to the CRA age-71+ schedule '
                f'this year.\n'
                f'  \u2022  New mandatory minimum: {fmt_pct(min_pct)} '
                f'\u00d7 {fmt_dollar(rrsp_bal)} = {fmt_dollar(min_amt)} '
                f'(up from the 1/(90\u2212age) formula used before 71).'
                + (f'\n  \u2022  Sarah continues to use Ralph\'s age ({spouse_age}) '
                   f'for the minimum rate ({fmt_pct(min_pct)} vs '
                   f'{fmt_pct(rrif_min_rate(age))} at own age).'
                   if use_spouse_age else '')
            ))
        else:
            spouse_note = (
                f'  \u2022  Sarah has elected to use Ralph\'s age ({spouse_age}) for the '
                f'RRIF minimum — rate reduced from '
                f'{fmt_pct(rrif_min_rate(age))} (own age) to '
                f'{fmt_pct(min_pct)} (Ralph\'s age).  '
                f'This election must be made at the time the RRIF is opened and '
                f'cannot be changed later.\n'
            ) if use_spouse_age else ''
            events.append((
                f'{name} \u2014 Convert RRSP to RRIF by Dec 31 (Mandatory)',
                f'{name}\'s RRSP (current balance: {fmt_dollar(rrsp_bal)}) must be '
                f'converted to a RRIF before December 31 of this year.  '
                f'No tax is triggered by the conversion.\n'
                f'{spouse_note}'
                f'  \u2022  Contact your financial institution before year-end.\n'
                f'  \u2022  First mandatory minimum: '
                f'{fmt_pct(min_pct)} \u00d7 {fmt_dollar(rrsp_bal)} = {fmt_dollar(min_amt)} '
                f'(due in the calendar year AFTER conversion).'
            ))

    if not events:
        return False, ch_rrif_open, li_rrif_open, ch_lif_open, li_lif_open

    add_sub_label(doc, 'Account Changes This Year \u2014 Action Required',
                  0x7B, 0x3F, 0x00, space_before=0)

    # Two-column table: title (narrow) | detail (wide)
    t = make_table(doc, 1 + len(events), [2.2, 4.3], border_color='E65100')

    hdr = t.rows[0].cells
    write_hdr_cell(hdr[0], 'Event',  CLR_EVENT_HDR)
    write_hdr_cell(hdr[1], 'What to Do', CLR_EVENT_HDR)

    for i, (title, detail) in enumerate(events):
        bg = CLR_EVENT_ROW
        cells = t.rows[i + 1].cells
        set_cell_bg(cells[0], bg); set_cell_pad(cells[0])
        set_cell_bg(cells[1], bg); set_cell_pad(cells[1])

        # Title cell: bold, amber-coloured text
        cells[0].text = ''
        p0 = cells[0].paragraphs[0]
        r0 = p0.add_run(title.replace('\u26a0\ufe0f  ', ''))
        r0.bold = True; r0.font.size = Pt(8); r0.font.name = 'Arial'
        r0.font.color.rgb = rgb('7B3F00')

        # Detail cell: multi-line, bullet lines rendered individually
        cells[1].text = ''
        first_para = True
        for line in detail.split('\n'):
            line = line.strip()
            if not line:
                continue
            if first_para:
                p1 = cells[1].paragraphs[0]
                first_para = False
            else:
                p1 = cells[1].add_paragraph()
            p1.paragraph_format.space_before = Pt(0)
            p1.paragraph_format.space_after  = Pt(1)
            r1 = p1.add_run(line)
            r1.font.size = Pt(8); r1.font.name = 'Arial'

    return True, ch_rrif_open, li_rrif_open, ch_lif_open, li_lif_open


# ---------------------------------------------------------------------------
# 1. Account Withdrawals
# ---------------------------------------------------------------------------

def build_withdrawals_table(doc, row):
    ch_age = int(fv(row, C_CH_AGE))
    li_age = int(fv(row, C_LI_AGE))

    items = []

    # Clamp to 0: negative values signal an RRSP *contribution* (shown in
    # Portfolio Management); they must not appear as withdrawals here.
    c_v = max(0.0, fv(row, C_CH_RRSP)); l_v = max(0.0, fv(row, C_LI_RRSP))
    if c_v > 0 or l_v > 0:
        rrif = (ch_age >= RRIF_CONVERSION_AGE + 1 or li_age >= RRIF_CONVERSION_AGE + 1)
        items.append(('RRIF Minimum Withdrawal' if rrif else 'RRSP / RRIF Withdrawal',
                       c_v, l_v,
                       'Minimum required by CRA' if rrif else ''))

    c_v = fv(row, C_CH_DCPP); l_v = fv(row, C_LI_DCPP)
    if c_v > 0 or l_v > 0:
        items.append(('DCPP / LIF Withdrawal', c_v, l_v, 'Annual minimum required'))

    c_v = fv(row, C_CH_TFSA); l_v = fv(row, C_LI_TFSA)
    if c_v > 0 or l_v > 0:
        items.append(('TFSA Withdrawal', c_v, l_v, 'Tax-free; no impact on income'))

    c_v = fv(row, C_CH_NONREG); l_v = fv(row, C_LI_NONREG)
    if c_v > 0 or l_v > 0:
        items.append(('Non-Registered Withdrawal', c_v, l_v, 'Capital gains may apply'))

    if not items:
        p = doc.add_paragraph('No account withdrawals required this year.')
        p.runs[0].font.size = Pt(9); p.runs[0].italic = True
        return

    c_tot = sum(x[1] for x in items)
    l_tot = sum(x[2] for x in items)
    t = make_table(doc, 1 + len(items) + 1, W4)

    hdr = t.rows[0].cells
    for cell, text in zip(hdr, ['Account / Action', 'Ralph', 'Sarah', 'Combined']):
        write_hdr_cell(cell, text, CLR_WITHDRAW_HDR,
                       align='center' if text != 'Account / Action' else 'left')

    for i, (label, c_v, l_v, note) in enumerate(items):
        bg = CLR_ROW_EVEN if i % 2 == 0 else CLR_WHITE
        cells = t.rows[i + 1].cells
        disp = label + (f'  \u2014  {note}' if note else '')
        write_data_cell(cells[0], disp, bg=bg)
        write_data_cell(cells[1], fmt_dollar(c_v), bg=bg, align='right')
        write_data_cell(cells[2], fmt_dollar(l_v), bg=bg, align='right')
        write_data_cell(cells[3], fmt_dollar(c_v + l_v), bg=bg, align='right')

    tot = t.rows[-1].cells
    for cell in tot:
        set_cell_bg(cell, CLR_TOTAL_WD); set_cell_pad(cell)
    write_cell(tot[0], 'Total to Withdraw', bold=True, size=9)
    write_cell(tot[1], fmt_dollar(c_tot), bold=True, align='right', size=9)
    write_cell(tot[2], fmt_dollar(l_tot), bold=True, align='right', size=9)
    write_cell(tot[3], fmt_dollar(c_tot + l_tot), bold=True, align='right', size=9)

# ---------------------------------------------------------------------------
# 2. Expected Income
# ---------------------------------------------------------------------------

def build_income_table(doc, row, first_cpp, first_oas):
    items = []

    c_v = fv(row, C_CH_SALARY); l_v = fv(row, C_LI_SALARY)
    if c_v > 0 or l_v > 0:
        items.append(('Employment Salary', c_v, l_v, ''))

    c_v = fv(row, C_CH_CPP); l_v = fv(row, C_LI_CPP)
    if c_v > 0 or l_v > 0:
        items.append(('Canada Pension Plan (CPP)', c_v, l_v,
                       'Begins this year \u2014 apply at Service Canada' if first_cpp else ''))

    c_v = fv(row, C_CH_OAS); l_v = fv(row, C_LI_OAS)
    if c_v > 0 or l_v > 0:
        items.append(('Old Age Security (OAS)', c_v, l_v,
                       'Begins this year \u2014 apply at Service Canada' if first_oas else ''))

    c_v = fv(row, C_CH_RENT); l_v = fv(row, C_LI_RENT)
    if c_v > 0 or l_v > 0:
        items.append(('Moonbrook Street Rental Income', c_v, l_v, ''))

    c_v = fv(row, C_CH_POLARON); l_v = fv(row, C_LI_POLARON)
    if c_v > 0 or l_v > 0:
        items.append(('Polaron Solar Income', c_v, l_v, ''))

    if not items:
        p = doc.add_paragraph('No automatic income sources this year.')
        p.runs[0].font.size = Pt(9); p.runs[0].italic = True
        return

    c_tot = sum(x[1] for x in items)
    l_tot = sum(x[2] for x in items)
    t = make_table(doc, 1 + len(items) + 1, W4)

    hdr = t.rows[0].cells
    for cell, text in zip(hdr, ['Income Source', 'Ralph', 'Sarah', 'Combined']):
        write_hdr_cell(cell, text, CLR_INCOME_HDR,
                       align='center' if text != 'Income Source' else 'left')

    for i, (label, c_v, l_v, note) in enumerate(items):
        bg = CLR_ROW_EVEN if i % 2 == 0 else CLR_WHITE
        cells = t.rows[i + 1].cells
        disp = label + (f'  \u2014  {note}' if note else '')
        write_data_cell(cells[0], disp, bg=bg)
        write_data_cell(cells[1], fmt_dollar(c_v), bg=bg, align='right')
        write_data_cell(cells[2], fmt_dollar(l_v), bg=bg, align='right')
        write_data_cell(cells[3], fmt_dollar(c_v + l_v), bg=bg, align='right')

    tot = t.rows[-1].cells
    for cell in tot:
        set_cell_bg(cell, CLR_TOTAL_INC); set_cell_pad(cell)
    write_cell(tot[0], 'Total Income', bold=True, size=9)
    write_cell(tot[1], fmt_dollar(c_tot), bold=True, align='right', size=9)
    write_cell(tot[2], fmt_dollar(l_tot), bold=True, align='right', size=9)
    write_cell(tot[3], fmt_dollar(c_tot + l_tot), bold=True, align='right', size=9)

# ---------------------------------------------------------------------------
# 3. Portfolio Management  (investments + non-reg tax harvesting)
# ---------------------------------------------------------------------------

def build_portfolio_table(doc, row):
    items = []

    # ---- RRSP contributions (working years / moonbrook sale year) ----
    # withdrawal_amount is negative when the simulation made a contribution.
    ch_rrsp_raw = fv(row, C_CH_RRSP)
    li_rrsp_raw = fv(row, C_LI_RRSP)
    ch_contrib = -ch_rrsp_raw if ch_rrsp_raw < 0 else 0.0
    li_contrib = -li_rrsp_raw if li_rrsp_raw < 0 else 0.0
    if ch_contrib > 0 or li_contrib > 0:
        ch_room_after = fv(row, A_CH_RRSP_ROOM)
        li_room_after = fv(row, A_LI_RRSP_ROOM)
        year = int(fv(row, C_YEAR))
        ch_salary = fv(row, C_CH_SALARY)
        li_salary = fv(row, C_LI_SALARY)

        def _room_str(amount):
            return fmt_dollar(amount) if amount > 0 else 'no room remaining'

        room_note = []
        if ch_contrib > 0:
            room_note.append(f'Ralph room remaining: {_room_str(ch_room_after)}')
        if li_contrib > 0:
            room_note.append(f'Sarah room remaining: {_room_str(li_room_after)}')

        # Explain contributions that appear in a year with no salary
        # (e.g. the 2035 Moonbrook-sale windfall year): the sim is absorbing
        # previously-unused carry-forward room to shelter the capital gain.
        post_retirement_note = ''
        if (ch_salary == 0 and li_salary == 0 and
            year > max(RALPH_RETIREMENT_YEAR, SARAH_RETIREMENT_YEAR)):
            post_retirement_note = (
                '  Note: this contribution uses carry-forward room from prior '
                'working years to shelter a large capital-gain event; no new '
                'room is being earned (no salary).'
            )

        items.append((
            'RRSP Contribution',
            ch_contrib, li_contrib,
            'Deposit to RRSP before Dec 31 — reduces taxable income this year. '
            + '  '.join(room_note)
            + post_retirement_note
        ))

    # ---- TFSA room remaining ----
    ch_tfsa_room = fv(row, A_CH_TFSA_ROOM)
    li_tfsa_room = fv(row, A_LI_TFSA_ROOM)
    if ch_tfsa_room > 0 or li_tfsa_room > 0:
        items.append(('TFSA Contribution Room Remaining',
                       ch_tfsa_room, li_tfsa_room,
                       'Can invest surplus cash here; withdrawals replenish room next Jan 1'))

    # ---- Non-reg portfolio balance ----
    ch_nreg = fv(row, A_CH_NREG_VAL)
    li_nreg = fv(row, A_LI_NREG_VAL)
    if ch_nreg > 0 or li_nreg > 0:
        items.append(('Non-Registered Portfolio Value',
                       ch_nreg, li_nreg, ''))

    # ---- Non-reg harvest: two separate actionable rows ----
    ch_sell = fv(row, C_CH_HARVEST)
    li_sell = fv(row, C_LI_HARVEST)
    ch_cap  = fv(row, C_CH_CAPITAL)
    li_cap  = fv(row, C_LI_CAPITAL)
    if ch_sell > 0 or li_sell > 0:
        items.append(('Non-Reg Harvest \u2014 Sell & Repurchase',
                       ch_sell, li_sell,
                       'Sell these securities then immediately rebuy at same price to reset ACB.  Note: the 30-day superficial-loss rule applies only to crystallised losses, NOT to harvested gains — same-day repurchase is fine here.'))
        items.append(('  Capital Gain Triggered (50% taxable)',
                       ch_cap, li_cap,
                       'Included in income above; no additional cash required'))

    if not items:
        p = doc.add_paragraph('No portfolio actions identified for this year.')
        p.runs[0].font.size = Pt(9); p.runs[0].italic = True
        return

    t = make_table(doc, 1 + len(items), W4)

    hdr = t.rows[0].cells
    for cell, text in zip(hdr, ['Portfolio Item', 'Ralph', 'Sarah', 'Notes']):
        write_hdr_cell(cell, text, CLR_PORTF_HDR,
                       align='center' if text not in ('Portfolio Item', 'Notes') else 'left')

    for i, (label, c_v, l_v, note) in enumerate(items):
        bg = CLR_ROW_EVEN if i % 2 == 0 else CLR_WHITE
        cells = t.rows[i + 1].cells
        write_data_cell(cells[0], label, bg=bg)
        write_data_cell(cells[1], fmt_dollar(c_v), bg=bg, align='right')
        write_data_cell(cells[2], fmt_dollar(l_v), bg=bg, align='right')
        write_data_cell(cells[3], note, bg=bg, italic=True, size=8)

# ---------------------------------------------------------------------------
# 4. Tax Summary
# ---------------------------------------------------------------------------

def build_tax_table(doc, row):
    # Eligible pension income = RRSP/RRIF withdrawals + LIF withdrawals.
    # This is the amount potentially eligible for T1032 pension splitting
    # (subject to the donor spouse being age 65+).
    # Negative RRSP values mean a contribution was made that year — clamp to
    # zero so a contribution year doesn't show phantom eligible pension income.
    ch_eligible = max(0.0, fv(row, C_CH_RRSP)) + max(0.0, fv(row, C_CH_DCPP))
    li_eligible = max(0.0, fv(row, C_LI_RRSP)) + max(0.0, fv(row, C_LI_DCPP))

    rows_data = [
        ('Gross Income',              fv(row, C_CH_INCOME),  fv(row, C_LI_INCOME),  False),
        ('  of which: RRIF/LIF Income (pension credit & T1032 eligible at 65)',
                                      ch_eligible,            li_eligible,            False),
        ('  of which: Capital Gains',  fv(row, C_CH_CAPITAL), fv(row, C_LI_CAPITAL), False),
        ('Tax Payable',               fv(row, C_CH_TAX),     fv(row, C_LI_TAX),     True),
        ('Effective Tax Rate',        None,                   None,                   False),
        ('After-Tax Income',          fv(row, C_CH_TAKE),    fv(row, C_LI_TAKE),    True),
    ]

    t = make_table(doc, 1 + len(rows_data), W3)

    hdr = t.rows[0].cells
    for cell, text in zip(hdr, ['Tax Summary', 'Ralph', 'Sarah']):
        write_hdr_cell(cell, text, CLR_TAX_HDR,
                       align='center' if text != 'Tax Summary' else 'left')

    for i, (label, c_v, l_v, bold) in enumerate(rows_data):
        bg = CLR_ROW_EVEN if i % 2 == 0 else CLR_WHITE
        cells = t.rows[i + 1].cells
        if label == 'Effective Tax Rate':
            c_str = fmt_pct(fv(row, C_CH_RATE))
            l_str = fmt_pct(fv(row, C_LI_RATE))
        else:
            c_str = fmt_dollar(c_v)
            l_str = fmt_dollar(l_v)
        write_data_cell(cells[0], label,  bg=bg, bold=bold)
        write_data_cell(cells[1], c_str,  bg=bg, bold=bold, align='right')
        write_data_cell(cells[2], l_str,  bg=bg, bold=bold, align='right')

# ---------------------------------------------------------------------------
# Pension-split optimizer helpers
# ---------------------------------------------------------------------------

def _marginal_rate(income, year, receives_oas):
    """
    Approximate combined federal + Ontario marginal rate for an Ontario
    resident with the given total income in the given calendar year.
    Includes the OAS clawback surcharge (15 pp) when the income exceeds the
    indexed threshold and the person receives OAS.
    """
    adj = (1.0 + _INFLATION) ** (year - _CURRENT_YEAR)

    fed = _FED_RATES[-1]
    for thresh, rate in zip(_FED_THRESHOLDS_2026, _FED_RATES):
        if income <= thresh * adj:
            fed = rate
            break

    ont = _ONT_RATES[-1]
    for thresh, rate in zip(_ONT_THRESHOLDS_2026, _ONT_RATES):
        if income <= thresh * adj:
            ont = rate
            break

    clawback = 0.15 if (receives_oas and
                        income > OAS_CLAWBACK_THRESHOLD * adj) else 0.0
    return fed + ont + clawback


def _optimal_pension_split(donor_pre, recip_pre, max_split, year,
                           donor_oas, recip_oas):
    """
    Binary-search for the T1032 split amount s in [0, max_split] that
    minimises combined tax by equalising marginal rates.

    The tax-minimising split satisfies:
        marginal_rate(donor_pre - s) == marginal_rate(recip_pre + s)

    Because both rate functions are non-decreasing step functions, the
    search keeps increasing s while the donor's rate exceeds the
    recipient's, then stops.  48 iterations gives < $1 precision on any
    realistic split range.

    Returns 0.0 if the donor's rate is already <= the recipient's
    (no split is beneficial).
    """
    if _marginal_rate(donor_pre, year, donor_oas) <= \
       _marginal_rate(recip_pre, year, recip_oas):
        return 0.0

    lo, hi = 0.0, max_split
    for _ in range(48):
        mid = (lo + hi) * 0.5
        if _marginal_rate(donor_pre - mid, year, donor_oas) > \
           _marginal_rate(recip_pre + mid, year, recip_oas):
            lo = mid    # still beneficial: donor rate still higher
        else:
            hi = mid    # split too far: recipient rate now higher

    # Ceiling-round to the nearest $500 so the recommendation clears any
    # bracket boundary cleanly and the displayed post-split rates are correct.
    # Clamp so rounding never exceeds the T1032 cap.
    raw = (lo + hi) * 0.5
    return min(max_split, math.ceil(raw / 500) * 500)


# ---------------------------------------------------------------------------
# 5. Tax Filing Notes  (conditional checklist)
# ---------------------------------------------------------------------------

def build_tax_notes(doc, row, first_cpp, first_oas, moonbrook_sale_year,
                    prev_ch_cpp, prev_li_cpp, prev_ch_oas, prev_li_oas,
                    prev_ch_tax=0.0, prev_li_tax=0.0):
    year   = int(fv(row, C_YEAR))
    ch_age = int(fv(row, C_CH_AGE))
    li_age = int(fv(row, C_LI_AGE))

    ch_income = fv(row, C_CH_INCOME)
    li_income = fv(row, C_LI_INCOME)
    ch_reg    = fv(row, C_CH_REG)
    li_reg    = fv(row, C_LI_REG)
    ch_cap    = fv(row, C_CH_CAPITAL)
    li_cap    = fv(row, C_LI_CAPITAL)
    ch_oas    = fv(row, C_CH_OAS)
    li_oas    = fv(row, C_LI_OAS)

    # Eligible pension income for T1032 = RRSP/RRIF + LIF withdrawals per spouse.
    # (These are the actual account debits before any pension-split adjustment.)
    # Clamp to zero: negative RRSP values mean a contribution, not a withdrawal.
    ch_eligible_pension = max(0.0, fv(row, C_CH_RRSP)) + max(0.0, fv(row, C_CH_DCPP))
    li_eligible_pension = max(0.0, fv(row, C_LI_RRSP)) + max(0.0, fv(row, C_LI_DCPP))

    # Indexed clawback threshold — mirrors tax.c:get_oas_clawback_threshold().
    # Base is CURRENT_YEAR (not 2024); rate is INFLATION (2.5%, not 2%).
    clawback_thresh = OAS_CLAWBACK_THRESHOLD * ((1.0 + _INFLATION) ** (year - _CURRENT_YEAR))

    notes = []   # list of (item_text, detail_text)

    # ---- Pension income splitting (T1032) ----
    # RRIF and LIF withdrawals only qualify for pension splitting once the
    # *transferring* spouse turns 65.  Before that age, neither RRSP withdrawals
    # nor LIF payments qualify (only a true RPP annuity would, which we don't
    # model).  The simulation begins splitting at PENSION_SPLIT_AGE = 65.
    older_age  = max(ch_age, li_age)
    if older_age >= PENSION_CREDIT_AGE and (ch_eligible_pension > 0 or li_eligible_pension > 0):
        # Identify donor (higher eligible pension) and recipient.
        if ch_eligible_pension >= li_eligible_pension:
            donor_name, recip_name = 'Ralph', 'Sarah'
            donor_age  = ch_age
            donor_elig = ch_eligible_pension
            recip_elig = li_eligible_pension
        else:
            donor_name, recip_name = 'Sarah', 'Ralph'
            donor_age  = li_age
            donor_elig = li_eligible_pension
            recip_elig = ch_eligible_pension

        if donor_age >= PENSION_CREDIT_AGE:
            max_split = donor_elig * 0.50

            # Reconstruct each spouse's pre-split income from individual
            # income-source columns.  C_CH_INCOME / C_LI_INCOME are POST the
            # sim's internal pension split, so using them directly produces a
            # near-zero gap and falls back to a flat 25% heuristic.  Summing
            # the raw component columns gives the true pre-split baseline.
            ch_pre = (ch_eligible_pension
                      + fv(row, C_CH_SALARY) + fv(row, C_CH_CPP)
                      + fv(row, C_CH_OAS)    + fv(row, C_CH_RENT)
                      + fv(row, C_CH_POLARON)+ fv(row, C_CH_CAPITAL))
            li_pre = (li_eligible_pension
                      + fv(row, C_LI_SALARY) + fv(row, C_LI_CPP)
                      + fv(row, C_LI_OAS)    + fv(row, C_LI_RENT)
                      + fv(row, C_LI_POLARON)+ fv(row, C_LI_CAPITAL))

            donor_pre  = ch_pre if donor_name == 'Ralph' else li_pre
            recip_pre  = li_pre if donor_name == 'Ralph' else ch_pre
            donor_oas  = fv(row, C_CH_OAS if donor_name == 'Ralph' else C_LI_OAS) > 0
            recip_oas  = fv(row, C_LI_OAS if donor_name == 'Ralph' else C_CH_OAS) > 0

            split_amt  = _optimal_pension_split(
                donor_pre, recip_pre, max_split, year, donor_oas, recip_oas)

            rate_d = _marginal_rate(donor_pre - split_amt, year, donor_oas)
            rate_r = _marginal_rate(recip_pre + split_amt, year, recip_oas)

            if split_amt < 100:
                split_detail = (
                    f'No T1032 split recommended — marginal rates already '
                    f'equalized at {rate_d:.0%}.  '
                    f'Max transferable if desired: {fmt_dollar(max_split)}.'
                )
            else:
                split_detail = (
                    f'Optimal split to {recip_name}: ~{fmt_dollar(split_amt)} '
                    f'(marginal rates after split — '
                    f'{donor_name}: {rate_d:.0%}, {recip_name}: {rate_r:.0%}; '
                    f'max: {fmt_dollar(max_split)}).  '
                    f'Both spouses must sign Form T1032 and attach to their returns.'
                )

            notes.append((
                'T1032 \u2014 Pension Income Splitting',
                f'{donor_name} is the transferring spouse.  '
                f'Eligible pension income: {fmt_dollar(donor_elig)} '
                f'(RRSP/RRIF {fmt_dollar(fv(row, C_CH_RRSP if donor_name == "Ralph" else C_LI_RRSP))} '
                f'+ LIF {fmt_dollar(fv(row, C_CH_DCPP if donor_name == "Ralph" else C_LI_DCPP))}).  '
                + split_detail
            ))
        else:
            # Younger spouse (donor by eligible income) is under 65 — can't split yet.
            notes.append((
                'T1032 \u2014 Pension Income Splitting Not Yet Available',
                f'{donor_name} holds more eligible pension income ({fmt_dollar(donor_elig)}) '
                f'but is {donor_age}, under the required age of 65.  '
                f'RRIF and LIF withdrawals do not qualify until age 65.  '
                f'Pension splitting becomes available to {donor_name} in '
                f'{year + (PENSION_CREDIT_AGE - donor_age)}.'
            ))

    # ---- Pension income credit ----
    # Uses post-split eligible pension amounts so that a spouse who received
    # split income (and now has >= $2,000) can also claim the credit.
    ch_pension_for_credit = ch_eligible_pension
    li_pension_for_credit = li_eligible_pension
    # If T1032 splitting occurred, the recipient gains eligible income.
    # Approximate: if incomes are equal and both have some eligible amount, credit both.
    credit_spouses = []
    if ch_age >= PENSION_CREDIT_AGE and ch_pension_for_credit > 0:
        credit_spouses.append(f'Ralph ({fmt_dollar(min(ch_pension_for_credit, PENSION_CREDIT_MIN))} credit base)')
    elif ch_age >= PENSION_CREDIT_AGE and li_age >= PENSION_CREDIT_AGE and li_eligible_pension > 0:
        # Ralph received split income from Sarah; ensure at least $2k attributed
        credit_spouses.append(f'Ralph (verify ≥{fmt_dollar(PENSION_CREDIT_MIN)} attributed via T1032)')
    if li_age >= PENSION_CREDIT_AGE and li_pension_for_credit > 0:
        credit_spouses.append(f'Sarah ({fmt_dollar(min(li_pension_for_credit, PENSION_CREDIT_MIN))} credit base)')
    if credit_spouses:
        notes.append((
            'Pension Income Credit \u2014 Line 31400',
            f'Claim for: {", ".join(credit_spouses)}.  '
            f'Worth up to $401/person (federal + Ontario combined) on first '
            f'{fmt_dollar(PENSION_CREDIT_MIN)} of eligible pension income.'
        ))

    # ---- CPP registration ----
    if first_cpp:
        who = []
        if fv(row, C_CH_CPP) > 0 and prev_ch_cpp == 0:
            who.append('Ralph')
        if fv(row, C_LI_CPP) > 0 and prev_li_cpp == 0:
            who.append('Sarah')
        if who:
            notes.append((
                f'CPP Retirement Benefits Begin \u2014 {", ".join(who)}',
                'Apply online at My Service Canada Account or in person. '
                'Allow 6\u201312 weeks for processing. '
                'CPP assignment between spouses may reduce combined tax.'
            ))

    # ---- OAS registration ----
    if first_oas:
        who = []
        if ch_oas > 0 and prev_ch_oas == 0:
            who.append('Ralph')
        if li_oas > 0 and prev_li_oas == 0:
            who.append('Sarah')
        if who:
            notes.append((
                f'OAS Begins \u2014 {", ".join(who)}',
                'OAS payments begin this year.  Confirm your Service Canada application '
                'was submitted and direct-deposit details are correct.  '
                'Check your My Service Canada Account for the first payment date.'
            ))

    # ---- OAS clawback watch ----
    for name, income, oas in (('Ralph', ch_income, ch_oas),
                               ('Sarah',  li_income, li_oas)):
        if oas > 0 and income > clawback_thresh * 0.90:
            notes.append((
                f'OAS Clawback Risk \u2014 {name}',
                f'{name}\u2019s net income ({fmt_dollar(income)}) is approaching the '
                f'clawback threshold ({fmt_dollar(clawback_thresh)}).  '
                'OAS is recovered at 15\u00a2 per dollar above threshold.  '
                'Shift income to TFSA withdrawals where possible. File Schedule T1206.'
            ))

    # ---- RRSP \u2192 RRIF conversion ----
    for name, age in (('Ralph', ch_age), ('Sarah', li_age)):
        if age == RRIF_CONVERSION_AGE:
            notes.append((
                f'RRSP \u2192 RRIF Conversion Required \u2014 {name}',
                f'{name} turns {age} this year.  RRSP must be converted to a RRIF '
                f'(or annuity) by December 31.  '
                'Elect the younger spouse as the RRIF annuitant to use their lower '
                'minimum withdrawal rates and defer tax.'
            ))

    # ---- Capital gains on Schedule 3 ----
    if ch_cap > 0 or li_cap > 0:
        cg_parts = ['Report all dispositions on Schedule 3.']
        if ch_cap > 0:
            cg_parts.append(
                f'Ralph: {fmt_dollar(ch_cap)} gain '
                f'({fmt_dollar(ch_cap * 0.5)} taxable at 50% inclusion).')
        if li_cap > 0:
            cg_parts.append(
                f'Sarah: {fmt_dollar(li_cap)} gain '
                f'({fmt_dollar(li_cap * 0.5)} taxable at 50% inclusion).')
        cg_parts.append(
            'If non-reg securities were sold and repurchased for ACB reset, '
            'document each transaction with date, proceeds, and adjusted cost base.')
        notes.append((
            'Capital Gains \u2014 Schedule 3',
            '  '.join(cg_parts)
        ))

    # ---- Moonbrook Street sale year ----
    if moonbrook_sale_year and year == moonbrook_sale_year:
        notes.append((
            'Moonbrook Street Property Sale \u2014 Critical Year',
            'Report disposition on Schedule 3. '
            'Capital gain = net proceeds \u2212 adjusted cost base. '
            'CCA recapture (undepreciated capital cost shortfall) is fully taxable as income \u2014 '
            'report on T776 / T2125. '
            'Consider RRSP/RRIF withdrawals minimal this year; '
            'capital-gain income will push marginal rates high. '
            'HST may apply on the commercial portion of the sale.'
        ))

    # ---- Per-spouse tax values used by several sections below ----
    ch_tax = fv(row, C_CH_TAX)
    li_tax = fv(row, C_LI_TAX)
    combined_tax = ch_tax + li_tax

    # ---- Year after Moonbrook Street sale: elect current-year installment basis ----
    # In the year following the sale, CRA\u2019s default installment demand is based on
    # the prior year\u2019s (enormous) tax bill.  Electing the current-year basis
    # keeps quarterly payments tied to this year\u2019s much-lower liability.
    if moonbrook_sale_year and year == moonbrook_sale_year + 1:
        cy_combined   = ch_tax + li_tax          # current year\u2019s total tax
        py_combined   = prev_ch_tax + prev_li_tax # prior year (sale year) total tax
        cy_quarterly  = cy_combined / 4.0
        # Per-spouse split proportional to their individual taxes
        ch_share = ch_tax / cy_combined if cy_combined > 0 else 0.5
        li_share = li_tax / cy_combined if cy_combined > 0 else 0.5
        ch_quarterly  = cy_quarterly * ch_share
        li_quarterly  = cy_quarterly * li_share
        # Round up to nearest $250 per spouse as a small buffer
        import math as _math
        ch_rounded = _math.ceil(ch_quarterly / 250) * 250
        li_rounded = _math.ceil(li_quarterly / 250) * 250
        savings = py_combined - cy_combined

        detail = (
            f'Elect the current-year option for {year}. '
            f'When the first installment reminder arrives in February {year}, '
            f'call CRA (1\u2011800\u2011959\u20118281) or log into CRA My Account \u2192 Installments '
            f'and select \u201cCalculate on current-year basis.\u201d  '
            f'Remit \u00bc \u00d7 {fmt_dollar(cy_combined)} \u2248 {fmt_dollar(cy_quarterly)} per quarter combined '
            f'(roughly {fmt_dollar(ch_quarterly)} Ralph, {fmt_dollar(li_quarterly)} Sarah), '
            f'or round up to {fmt_dollar(ch_rounded)} Ralph / {fmt_dollar(li_rounded)} Sarah per quarter '
            f'as a small buffer.  '
            f'Total paid across four quarters: \u2248 {fmt_dollar(cy_combined)} combined, '
            f'vs. the {fmt_dollar(py_combined)} CRA would demand under the prior-year (no-calc) default.  '
            f'The difference \u2014 {fmt_dollar(savings)} \u2014 stays in your account rather than '
            f'sitting interest-free with CRA until your April {year + 1} refund.'
        ) if py_combined > cy_combined else (
            f'CRA will issue installment reminders in {year} based on last year\u2019s tax bill '
            f'({fmt_dollar(py_combined)}).  '
            f'Elect the current-year basis to limit quarterly payments to '
            f'\u00bc \u00d7 {fmt_dollar(cy_combined)} \u2248 {fmt_dollar(cy_quarterly)} per quarter combined.  '
            f'Call CRA (1\u2011800\u2011959\u20118281) or log into CRA My Account \u2192 Installments '
            f'to make the election before the first March instalment due date.'
        )
        notes.append((
            f'CRA Installment Election \u2014 Current-Year Basis',
            detail
        ))

    # ---- TFSA recontribution note ----
    ch_tfsa_wd = fv(row, C_CH_TFSA); li_tfsa_wd = fv(row, C_LI_TFSA)
    if ch_tfsa_wd > 0 or li_tfsa_wd > 0:
        notes.append((
            'TFSA Recontribution Room',
            'Amounts withdrawn this year are re-added to your TFSA room on '
            'January 1 of the following year.  '
            f'Ralph withdrew {fmt_dollar(ch_tfsa_wd)};  '
            f'Sarah withdrew {fmt_dollar(li_tfsa_wd)}.  '
            'Track room carefully to avoid the 1%/month over-contribution penalty.'
        ))

    # ---- Always: consider spousal RRSP if one spouse has much more in RRSP ----
    # Only relevant before RRIF conversion age AND while at least one spouse is
    # still working (contributions require earned income and available room;
    # post-retirement the suggestion no longer applies).
    both_retired = year > max(RALPH_RETIREMENT_YEAR, SARAH_RETIREMENT_YEAR)
    if not both_retired and ch_age < RRIF_CONVERSION_AGE and li_age < RRIF_CONVERSION_AGE:
        ch_rrsp_val = fv(row, A_CH_RRSP_VAL)
        li_rrsp_val = fv(row, A_LI_RRSP_VAL)
        if ch_rrsp_val > 0 and li_rrsp_val > 0:
            ratio = max(ch_rrsp_val, li_rrsp_val) / max(min(ch_rrsp_val, li_rrsp_val), 1)
            if ratio > 1.5:
                higher = 'Ralph' if ch_rrsp_val > li_rrsp_val else 'Sarah'
                notes.append((
                    'RRSP Balance Imbalance \u2014 Income Splitting Opportunity',
                    f'{higher} has a significantly larger RRSP balance.  '
                    'Consider whether contributing to a spousal RRSP is still available '
                    'to equalise future RRIF withdrawals and reduce combined tax.'
                ))

    # ---- Installment reminder triage ----
    # CRA sends installment reminders whenever "net tax owing" exceeds $3,000 in
    # the current year AND in either of the two preceding years.  Important:
    # "net tax owing" is line 43500 minus source withholding, NOT the balance
    # owing on the T1.  Voluntary lump-sum payments to CRA reduce balance owing
    # but do NOT reduce net tax owing for the threshold test.
    if combined_tax > 0:
        installment_threshold = 3000
        each_above = [
            f'Ralph {fmt_dollar(ch_tax)}' if ch_tax > installment_threshold else '',
            f'Sarah {fmt_dollar(li_tax)}'  if li_tax > installment_threshold else '',
        ]
        above_str = ',  '.join(x for x in each_above if x)
        if above_str:
            detail = (
                f'Estimated taxes this year: Ralph {fmt_dollar(ch_tax)},  '
                f'Sarah {fmt_dollar(li_tax)}.  '
                f'Both exceed the ${installment_threshold:,} threshold '
                f'({above_str}).  '
                if all(each_above) else
                f'Estimated taxes this year: Ralph {fmt_dollar(ch_tax)},  '
                f'Sarah {fmt_dollar(li_tax)}.  '
                f'{above_str} exceeds the ${installment_threshold:,} threshold.  '
            )
        else:
            detail = (
                f'Estimated taxes this year: Ralph {fmt_dollar(ch_tax)},  '
                f'Sarah {fmt_dollar(li_tax)} — both under the '
                f'${installment_threshold:,} threshold this year.  '
            )
        detail += (
            f'CRA tests "net tax owing" (line 43500 minus source withholding '
            f'and refundable credits), NOT the balance owing on the T1.  '
            f'If net tax owing exceeds ${installment_threshold:,} this year '
            f'AND exceeded ${installment_threshold:,} in either of the prior '
            f'two years, CRA will send installment reminders for next year.  '
            f'IMPORTANT: a voluntary lump-sum payment to CRA in December does '
            f'NOT reduce net tax owing for the threshold test \u2014 it only '
            f'reduces the final balance owing on the April return.  '
            f'Legitimate ways to avoid installments: '
            f'(a) when the installment reminder arrives, elect the '
            f'"current-year option" if actual current-year tax will be lower '
            f'(you pay based on your own estimate instead of CRA\u2019s); '
            f'(b) increase source withholding on RRIF / CPP / OAS by filing '
            f'Form T1213(OAS) with CRA or Form ISP-3520 with Service Canada '
            f'\u2014 source withholding DOES count against net tax owing; '
            f'(c) reduce taxable income via an RRSP contribution (if room '
            f'remains) or capital-loss harvesting before year-end.'
        )
        notes.append((
            f'Installment Reminders \u2014 How To Respond',
            detail
        ))

    # ---- Render ----
    if not notes:
        p = doc.add_paragraph('No specific tax filing actions identified for this year.')
        p.runs[0].font.size = Pt(9); p.runs[0].italic = True
        return

    t = make_table(doc, 1 + len(notes), W2)

    hdr = t.rows[0].cells
    write_hdr_cell(hdr[0], 'Item / Form', CLR_NOTES_HDR)
    write_hdr_cell(hdr[1], 'Action Required / Notes', CLR_NOTES_HDR)

    for i, (item, detail) in enumerate(notes):
        cells = t.rows[i + 1].cells
        set_cell_bg(cells[0], CLR_NOTES_ROW); set_cell_pad(cells[0])
        set_cell_bg(cells[1], CLR_WHITE);      set_cell_pad(cells[1])
        write_cell(cells[0], item,   bold=True,  size=8)
        _write_bulleted_detail(cells[1], detail)


def _write_bulleted_detail(cell, detail):
    """
    Render a long prose detail string as a short lead sentence plus
    bullet lines.  Splits the string on '.  ' (the standard separator
    used throughout build_tax_notes) so each sentence becomes its own
    paragraph — much easier to read in old age.  The first sentence
    remains un-bulleted as a lead-in; the rest become bullets.
    """
    cell.text = ''
    # Normalize multi-space separators used in our notes into a clean
    # sentence-break marker, then split.
    parts_raw = re.split(r'(?<=[.])\s{2,}', detail.strip())
    # Fall back to splitting on ".  " when no double-space separator exists.
    parts = [s.strip() for s in parts_raw if s.strip()]
    if len(parts) == 1:
        # Try splitting on ". " as a last resort
        parts = [s.strip() for s in re.split(r'(?<=[.])\s+', detail.strip())
                 if s.strip()]

    if not parts:
        write_cell(cell, detail, size=8)
        return

    # First sentence: plain paragraph (lead-in)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(1)
    r0 = p0.add_run(parts[0] if parts[0].endswith('.') else parts[0] + '.')
    r0.font.size = Pt(8); r0.font.name = 'Arial'

    # Remaining sentences: each prefixed with a bullet glyph and indented
    for part in parts[1:]:
        bp = cell.add_paragraph()
        bp.paragraph_format.left_indent  = Inches(0.12)
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after  = Pt(1)
        br = bp.add_run('\u2022  ' + (part if part.endswith('.') else part + '.'))
        br.font.size = Pt(8); br.font.name = 'Arial'

# ---------------------------------------------------------------------------
# 5b. Depletion Warnings  (accounts running low on balance vs. planned draw)
# ---------------------------------------------------------------------------

def build_depletion_warnings(doc, row):
    """
    Warn when an account balance is within roughly one year's withdrawal of
    being empty, or when a draw exceeds the current balance.  Helps the
    reader spot that an income stream is about to run out without having to
    eyeball the year-to-year asset columns.
    """
    warnings = []

    def _check(owner, label, balance, draw):
        if draw <= 0 or balance <= 0:
            return
        # balance is the END-of-year value (post-withdrawal, post-growth) as
        # written by output.c.  Do NOT subtract draw again — it has already
        # been deducted before the value was recorded.
        # Warn when the current balance supports fewer than ~2 more years at
        # the current draw rate, or when a single draw would wipe it out.
        years_left = balance / draw
        if years_left < 1.0:
            warnings.append((owner, label,
                f'Balance ({fmt_dollar(balance)}) is less than one year\'s draw '
                f'({fmt_dollar(draw)}).  This account may be exhausted next year; '
                f'plan an alternative income source now.'))
        elif years_left < 2.0:
            warnings.append((owner, label,
                f'Balance ({fmt_dollar(balance)}) covers roughly {years_left:.1f} '
                f'more year(s) at the current draw rate ({fmt_dollar(draw)}/yr).  '
                f'Plan an alternative income source before this account empties.'))

    _check('Ralph', 'RRSP / RRIF',    fv(row, A_CH_RRSP_VAL),  max(0.0, fv(row, C_CH_RRSP)))
    _check('Sarah',  'RRSP / RRIF',    fv(row, A_LI_RRSP_VAL),  max(0.0, fv(row, C_LI_RRSP)))
    _check('Ralph', 'DCPP / LIF',     fv(row, A_CH_DCPP_VAL),  fv(row, C_CH_DCPP))
    _check('Sarah',  'DCPP / LIF',     fv(row, A_LI_DCPP_VAL),  fv(row, C_LI_DCPP))
    _check('Ralph', 'TFSA',           fv(row, A_CH_TFSA_VAL),  fv(row, C_CH_TFSA))
    _check('Sarah',  'TFSA',           fv(row, A_LI_TFSA_VAL),  fv(row, C_LI_TFSA))
    _check('Ralph', 'Non-Registered', fv(row, A_CH_NREG_VAL),  fv(row, C_CH_NONREG))
    _check('Sarah',  'Non-Registered', fv(row, A_LI_NREG_VAL),  fv(row, C_LI_NONREG))

    if not warnings:
        return

    add_sub_label(doc, 'Depletion Warnings \u2014 Accounts Running Low',
                  0xC0, 0x00, 0x00, space_before=4)

    t = make_table(doc, 1 + len(warnings), [1.0, 1.6, 3.9], border_color='C00000')

    hdr = t.rows[0].cells
    write_hdr_cell(hdr[0], 'Who',     'C00000')
    write_hdr_cell(hdr[1], 'Account', 'C00000')
    write_hdr_cell(hdr[2], 'Warning', 'C00000')

    for i, (owner, label, msg) in enumerate(warnings):
        cells = t.rows[i + 1].cells
        bg = 'FFEBEE'
        set_cell_bg(cells[0], bg); set_cell_pad(cells[0])
        set_cell_bg(cells[1], bg); set_cell_pad(cells[1])
        set_cell_bg(cells[2], bg); set_cell_pad(cells[2])
        write_cell(cells[0], owner, bold=True, size=8, color='C00000')
        write_cell(cells[1], label, bold=True, size=8)
        write_cell(cells[2], msg, size=8)


# ---------------------------------------------------------------------------
# 6. Year-End Position
# ---------------------------------------------------------------------------

def build_networth_table(doc, row):
    cash = fv(row, C_CASH)
    nw   = fv(row, C_NETWORTH)
    exp  = abs(fv(row, C_EXPENSES))

    t = make_table(doc, 2, [1.65, 1.65, 1.65, 1.55])

    hdr = t.rows[0].cells
    for cell, text in zip(hdr, ['Annual Expenses',
                                'End-of-Year Cash',
                                '',
                                'Total Net Worth']):
        write_hdr_cell(cell, text, CLR_NW_HDR, align='center')

    dr = t.rows[1].cells
    for cell in dr:
        set_cell_bg(cell, CLR_NW_DATA); set_cell_pad(cell)
    write_cell(dr[0], fmt_dollar(exp, show_dash=False), bold=True, align='center', size=10)
    write_cell(dr[1], fmt_dollar(cash, show_dash=False), bold=True, align='center', size=10)
    write_cell(dr[2], '', size=9)
    write_cell(dr[3], fmt_dollar(nw,  show_dash=False), bold=True, align='center', size=10)

    # Footnote clarifying what Cash is and what Net Worth includes
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after  = Pt(0)
    nr = note.add_run(
        'End-of-Year Cash = leftover chequing balance after all withdrawals, '
        'income, spending, and taxes have cleared (rolls into next year).  '
        'Total Net Worth = sum of all accounts + real estate equity, '
        'after applying deemed-disposition taxes on deferred accounts.'
    )
    nr.italic = True; nr.font.size = Pt(8)
    nr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    in_csv  = sys.argv[1] if len(sys.argv) > 1 else 'retirement.csv'
    out_doc = sys.argv[2] if len(sys.argv) > 2 \
              else f"retirement_plan_{date.today().strftime('%Y-%m-%d')}.docx"

    # ---- Parse CSV ----
    # Read the main data block (stops at the first blank line) and then
    # continue to collect summary rows so we can find the Survivor metadata.
    all_rows     = []
    summary_rows = []
    with open(in_csv, newline='', encoding='utf-8-sig') as f:
        reader   = csv.reader(f, delimiter=';')
        in_summary = False
        for row in reader:
            if all(c.strip() == '' for c in row):
                in_summary = True
                continue
            if in_summary:
                summary_rows.append(row)
            else:
                all_rows.append(row)

    if len(all_rows) < 2:
        print(f'Error: no usable data in {in_csv}'); sys.exit(1)

    sim_rows = []
    for row in all_rows[1:]:
        try:
            if int(row[C_YEAR]) >= MIN_SIM_YEAR:
                sim_rows.append(row)
        except (ValueError, IndexError):
            pass

    if not sim_rows:
        print('No simulation rows found (expected year >= 2026).'); sys.exit(1)

    # ---- Parse survivor metadata ----
    _CH_AGE_COL = 26   # 0-based index of "Ralph Age" in data rows
    _LI_AGE_COL = 36   # 0-based index of "Sarah Age" in data rows
    survivor_name = None   # 'Ralph' or 'Sarah' if active, else None
    survivor_year = None
    survivor_age  = None
    for r in summary_rows:
        if r and r[0].strip() == "Survivor" and len(r) >= 3:
            name_raw = r[1].strip().lower()
            if name_raw != "none":
                survivor_name = name_raw.capitalize()
                try:
                    survivor_year = int(r[2].strip())
                except ValueError:
                    pass
                if survivor_year:
                    age_col = _CH_AGE_COL if survivor_name == "Ralph" else _LI_AGE_COL
                    for dr in sim_rows:
                        try:
                            if int(dr[C_YEAR]) == survivor_year and len(dr) > age_col:
                                survivor_age = int(float(dr[age_col]))
                                break
                        except (ValueError, IndexError):
                            pass
            break

    # ---- Parse realized-return averages ----
    # The simulator emits six summary rows (three pairs of Expected vs
    # Realized per asset class) so the document can show how the randomly-
    # drawn market path for this run compared to the long-run means.
    # Format: "<Expected|Realized> <Class> Return;;;<fraction>;"
    returns = {}    # e.g. ('Realized', 'Financial') -> 0.0779
    for r in summary_rows:
        if not r or len(r) < 4:
            continue
        label = r[0].strip()
        m = re.match(r'^(Expected|Realized|Min|Max)\s+(Financial|Property|Rent)\s+Return$', label)
        if not m:
            continue
        try:
            returns[(m.group(1), m.group(2))] = float(r[3].strip())
        except ValueError:
            pass

    # Determine the moonbrook sale year from the data (last year Moonbrook Rent income appears)
    moonbrook_sale_year = None
    for row in reversed(sim_rows):
        if fv(row, C_CH_RENT) > 0 or fv(row, C_LI_RENT) > 0:
            moonbrook_sale_year = int(fv(row, C_YEAR)) + 1
            break

    # Determine mortgage payoff years (year the balance first hits zero)
    _PB = _AB + 40
    A_CH_ROBERTSON_BAL = _PB + 5;  A_LI_ROBERTSON_BAL = _PB + 13
    A_CH_MOONBROOK_BAL   = _PB + 21; A_LI_MOONBROOK_BAL   = _PB + 29

    robertson_payoff_year = moonbrook_payoff_year = None
    prev_robertson_bal = prev_moonbrook_bal = None
    for row in sim_rows:
        try: year = int(row[C_YEAR])
        except: continue
        m_bal = fv(row, A_CH_ROBERTSON_BAL) + fv(row, A_LI_ROBERTSON_BAL)
        f_bal = fv(row, A_CH_MOONBROOK_BAL)   + fv(row, A_LI_MOONBROOK_BAL)
        if robertson_payoff_year is None and prev_robertson_bal is not None \
                and prev_robertson_bal > 0 and m_bal == 0:
            robertson_payoff_year = year
        if moonbrook_payoff_year is None and prev_moonbrook_bal is not None \
                and prev_moonbrook_bal > 0 and f_bal == 0:
            moonbrook_payoff_year = year
        prev_robertson_bal = m_bal
        prev_moonbrook_bal   = f_bal

    # Build milestone dict (year → list of strings)
    all_milestones = compute_all_milestones(
        moonbrook_sale_year, robertson_payoff_year, moonbrook_payoff_year,
        survivor_name=survivor_name,
        survivor_year=survivor_year,
        survivor_age=survivor_age)

    # ---- Build document ----
    doc = Document()

    for sec in doc.sections:
        sec.page_width    = Inches(8.5)
        sec.page_height   = Inches(11)
        sec.top_margin    = Inches(0.70)
        sec.bottom_margin = Inches(0.70)
        sec.left_margin   = Inches(0.75)
        sec.right_margin  = Inches(0.75)

    add_page_number_footer(doc)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    # ---- Cover page ----
    tp = doc.add_heading('Retirement Action Plan', 0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in tp.runs:
        run.font.color.rgb = RGBColor(*CLR_TITLE_BLUE)

    sub = doc.add_paragraph('Ralph & Sarah  \u2014  Year-by-Year Guide')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size      = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    dp = doc.add_paragraph(f'Generated {date.today().strftime("%B %d, %Y")}')
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dp.runs[0].italic         = True
    dp.runs[0].font.size      = Pt(10)
    dp.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()
    final = sim_rows[-1]
    summ  = doc.add_paragraph(
        f'Simulation covers {len(sim_rows)} years '
        f'({sim_rows[0][C_YEAR]}\u2013{final[C_YEAR]}).  '
        f'Projected final net worth: '
        f'{fmt_dollar(fv(final, C_NETWORTH), show_dash=False)}.'
    )
    summ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summ.runs[0].font.size = Pt(10)

    if moonbrook_sale_year:
        fs = doc.add_paragraph(
            f'Optimal Moonbrook Street sale year: {moonbrook_sale_year}.'
        )
        fs.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fs.runs[0].font.size = Pt(10)

    # ---- Strategy Robustness callout ----
    # Surfaces the eval-pass success ratio (how many of the
    # EVALUATION_PATHS_PER_STRATEGY re-test paths avoided cash exhaustion)
    # with a threshold classification.  Skipped silently if the CSV lacks
    # the row (older simulation binaries).
    successes, total = parse_success_ratio(summary_rows)
    if successes is not None and total:
        doc.add_paragraph()   # spacer before callout
        build_success_callout(doc, successes, total)

    # ---- Stochastic-returns note ----
    # Every trial draws a fresh market path — annual returns are sampled from
    # a normal distribution around each asset's expected rate.  We surface
    # both the base assumption and what actually occurred along the
    # representative path so the reader can judge how favourable / adverse
    # this specific sequence was.
    fin_vol = _p('FINANCIAL_RETURN_VOLATILITY', 0.0)
    pro_vol = _p('PROPERTY_RETURN_VOLATILITY',  0.0)
    ren_vol = _p('RENT_RETURN_VOLATILITY',      0.0)
    # Read actual run-time value from CSV (compiler flag may override params.h default).
    paths = None
    for _r in summary_rows:
        if _r and _r[0].strip() == 'Return Paths Per Strategy':
            for _cell in _r[1:]:
                s = _cell.strip()
                if s:
                    try:
                        paths = int(s)
                    except ValueError:
                        pass
                    break
    if paths is None:
        paths = _p('RETURN_PATHS_PER_STRATEGY', 1)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run(
        f'Returns are stochastic: each year, financial assets are drawn around '
        f'their expected rate with σ={fin_vol*100:.1f}%, property with '
        f'σ={pro_vol*100:.1f}%, and rent growth with σ={ren_vol*100:.1f}%.  '
        f'Each candidate withdrawal strategy is evaluated against '
        f'{int(paths)} independent market paths; the representative path '
        f'shown below is the one closest to its strategy\u2019s mean terminal '
        f'net worth.'
    )
    r.italic         = True
    r.font.size      = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if returns:
        def _pair(cls):
            exp = returns.get(('Expected', cls))
            rea = returns.get(('Realized', cls))
            if exp is None or rea is None:
                return None
            return f'{cls.lower()} {rea*100:+.2f}% vs {exp*100:+.2f}% expected'
        parts = [p for p in (_pair('Financial'),
                             _pair('Property'),
                             _pair('Rent')) if p]
        if parts:
            avg = doc.add_paragraph()
            avg.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ar = avg.add_run(
                'Representative path average returns: ' + '; '.join(parts) + '.'
            )
            ar.italic         = True
            ar.font.size      = Pt(9)
            ar.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        fin_min = returns.get(('Min', 'Financial'))
        fin_max = returns.get(('Max', 'Financial'))
        if fin_min is not None and fin_max is not None:
            rng = doc.add_paragraph()
            rng.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = rng.add_run(
                f'Financial return range: best year {fin_max*100:+.2f}%'
                f',  worst year {fin_min*100:+.2f}%.'
            )
            rr.italic         = True
            rr.font.size      = Pt(9)
            rr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ---- Intra-year withdrawal sequencing (rules of thumb) ----
    doc.add_paragraph()  # spacer
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hdr.add_run('Intra-Year Withdrawal Sequencing')
    hr.bold           = True
    hr.font.size      = Pt(10)
    hr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    bullets = [
        'Run month-to-month cash flow from the TFSA; leave the RRIF draw for '
        'December once full-year income is visible.',
        'Size the year-end RRIF top-up to land net income just below the OAS '
        'clawback threshold (~$95K in 2026, inflation-indexed) — and, earlier, below the Age Amount '
        'phase-out (~$44K in 2026, inflation-indexed) where feasible.',
        'Treat the TFSA as a shock absorber: cover unexpected expenses from '
        'there to avoid crossing a tax cliff late in the year.',
        'In a down market year, pull the RRIF before the TFSA — preserve '
        'tax-free compounding and shrink the mandatory base for next year.',
        'Above-minimum RRIF draws should come in December so withholding tax '
        'is refunded on the next return rather than sitting with CRA.',
    ]
    for b in bullets:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent  = Inches(0.75)
        p.paragraph_format.right_indent = Inches(0.75)
        p.paragraph_format.space_after  = Pt(2)
        br = p.add_run(b)
        br.font.size      = Pt(9)
        br.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ---- Front-matter aging-friendly reference pages ----
    # Keep these BEFORE the year-by-year body so they are easy to flip to
    # without getting lost in the 40+ year-pages that follow.
    build_table_of_contents(doc)
    build_rules_page(doc)
    build_red_flags_page(doc)
    build_contacts_page(doc)
    build_glossary_page(doc)
    build_standard_year_checklist(doc)
    build_surviving_spouse_page(doc)

    # ---- Milestone tracking ----
    cpp_seen = oas_seen = False
    prev_ch_cpp  = prev_li_cpp  = 0.0
    prev_ch_oas  = prev_li_oas  = 0.0
    prev_ch_rrsp = prev_li_rrsp = 0.0    # for early-RRIF detection
    prev_ch_tax  = prev_li_tax  = 0.0    # for moonbrook+1 installment election note
    ch_rrif_open = li_rrif_open = False   # set True once RRIF opened
    ch_lif_open  = li_lif_open  = False   # set True once LIF opened

    # ---- Per-year sections ----
    for idx, row in enumerate(sim_rows):
        if len(row) < 44:
            continue

        year   = row[C_YEAR].strip()
        ch_age = int(fv(row, C_CH_AGE))
        li_age = int(fv(row, C_LI_AGE))

        first_cpp = (not cpp_seen and
                     (fv(row, C_CH_CPP) > 0 or fv(row, C_LI_CPP) > 0))
        first_oas = (not oas_seen and
                     (fv(row, C_CH_OAS) > 0 or fv(row, C_LI_OAS) > 0))
        if first_cpp: cpp_seen = True
        if first_oas: oas_seen = True

        # ---- Heading ----
        h = doc.add_heading(
            f'{year}     Ralph: Age {ch_age}  |  Sarah: Age {li_age}', level=1)
        h.runs[0].font.size      = Pt(13)
        h.runs[0].font.color.rgb = RGBColor(*CLR_TITLE_BLUE)
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after  = Pt(2)

        # Life milestones banner (retirement, children's school, mortgages)
        build_milestones_banner(doc, all_milestones.get(int(year), []))

        # Life-transition callout for big decision years (retirement, Moonbrook
        # sale, age 65 CPP/OAS, age 70 RRIF approach, Robertson payoff, etc.)
        build_transition_callout(doc, int(year), ch_age, li_age,
                                 moonbrook_sale_year, robertson_payoff_year)

        # Account events: LIF opening, early RRIF, mandatory RRIF conversion
        _, ch_rrif_open, li_rrif_open, ch_lif_open, li_lif_open = build_account_events(
            doc, row,
            prev_ch_rrsp=prev_ch_rrsp, prev_li_rrsp=prev_li_rrsp,
            ch_rrif_open=ch_rrif_open,  li_rrif_open=li_rrif_open,
            ch_lif_open=ch_lif_open,    li_lif_open=li_lif_open,
        )

        add_sub_label(doc, 'Actions Required \u2014 Account Withdrawals',
                      0x1E, 0x56, 0x31, space_before=4)
        build_withdrawals_table(doc, row)

        add_sub_label(doc, 'Expected Income', 0x15, 0x65, 0xC0)
        build_income_table(doc, row, first_cpp, first_oas)

        add_sub_label(doc, 'Portfolio Management \u2014 Investments & Tax Harvesting',
                      0x4E, 0x34, 0x2E)
        build_portfolio_table(doc, row)

        add_sub_label(doc, 'Tax Summary', 0x4A, 0x23, 0x5A)
        build_tax_table(doc, row)

        add_sub_label(doc, 'Tax Filing Notes', 0x37, 0x47, 0x4F)
        build_tax_notes(doc, row, first_cpp, first_oas, moonbrook_sale_year,
                        prev_ch_cpp, prev_li_cpp, prev_ch_oas, prev_li_oas,
                        prev_ch_tax=prev_ch_tax, prev_li_tax=prev_li_tax)

        # Account depletion warnings (only emit when something is close to empty)
        build_depletion_warnings(doc, row)

        add_sub_label(doc, 'Year-End Position', 0x1F, 0x4E, 0x79)
        build_networth_table(doc, row)

        if idx < len(sim_rows) - 1:
            doc.add_page_break()

        # Track for next iteration's milestone detection
        prev_ch_cpp  = fv(row, C_CH_CPP)
        prev_li_cpp  = fv(row, C_LI_CPP)
        prev_ch_oas  = fv(row, C_CH_OAS)
        prev_li_oas  = fv(row, C_LI_OAS)
        prev_ch_rrsp = fv(row, C_CH_RRSP)
        prev_li_rrsp = fv(row, C_LI_RRSP)
        prev_ch_tax  = fv(row, C_CH_TAX)
        prev_li_tax  = fv(row, C_LI_TAX)

    doc.save(out_doc)
    print(f'Saved: {out_doc}  ({len(sim_rows)} years)')


if __name__ == '__main__':
    main()
